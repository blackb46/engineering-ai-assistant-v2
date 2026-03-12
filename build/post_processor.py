"""
post_processor.py
=================
City of Brentwood Engineering AI Assistant — V2
Pre-processing step that runs automatically inside build_corpus.py
before every document is chunked.

PURPOSE:
    Municipal code documents from Municode and the Engineering Policy
    Manual store numbered and bulleted content as "List Paragraph" style
    paragraphs. The section chunker reads these correctly, but long list
    sections get split across chunk boundaries — meaning critical facts
    (like driveway grade limits: 20% paved, 10% unpaved, 5% cross slope)
    land in two separate chunks where only one half scores high enough
    to retrieve.

    This module scans every document for contiguous groups of list items,
    joins them into readable prose, and inserts that prose as a Normal-style
    paragraph immediately after the list group. The original list items are
    untouched. The chunker then indexes both.

    Result: critical policy facts are reliably retrievable even when the
    original list items straddle a chunk boundary.

HOW IT WORKS:
    1. Load the DOCX into memory (original file never modified)
    2. Scan paragraphs in order, tracking the current section heading
    3. When a contiguous list group ends, build a prose summary of its items
    4. Insert the summary as a Normal paragraph after the last list item
    5. Save to a temporary file and return that path to build_corpus.py
    6. build_corpus.py passes the temp path to chunk_document() as normal
    7. Temp file is deleted by build_corpus.py after chunking

MAINTENANCE:
    You never need to edit source documents manually.
    When a code chapter is amended, just replace the DOCX in Google Drive
    and run the Colab builder — this module handles everything automatically.

FILE LOCATION: build/post_processor.py

NOTE: This file belongs ONLY in build/. Do not duplicate it in utils/.
      build_corpus.py imports from build/post_processor.py directly.
"""

import os
import re
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# STYLE CLASSIFICATION
# ─────────────────────────────────────────────────────────────────────────────

LIST_STYLES = {
    "List Paragraph", "ListParagraph",
    "List Bullet", "List Bullet 2", "List Bullet 3",
    "List Number", "List Number 2", "List Number 3",
    "List Continue", "List Continue 2",
    # Confirmed Municode DOCX styles from Chapter 78 diagnostic (March 2026):
    # style='list1'  — numbered items like "(1) Minimum required lot area..."
    # style='list2'  — sub-items like "(a) When the length of driveway exceeds..."
    # style='list3'  — sub-sub-items
    # style='list4'  — deepest nesting level
    "list1", "list2", "list3", "list4",
    # Additional Municode styles confirmed in corpus:
    # style='b2'   — body text continuation in some chapters
    # style='bc1'  — body continuation level 1
    # style='bc2'  — body continuation level 2
    # style='p0'   — introductory paragraph before a list group
    #               (e.g. "For all property within the R-2 zoning districts,
    #                the following minimum technical standards shall apply...")
    #               Including p0 ensures the intro sentence is captured in the
    #               prose summary alongside the list items that follow it.
    "b2", "bc1", "bc2",
    "p0",
    # Municode enumeration styles (alternative naming in some chapters)
    "enumeration1", "enumeration2", "enumeration3",
    "liststyle1", "liststyle2",
    "bodyindent", "bodytextindent",
}

SKIP_STYLES = {
    "historynote0", "refeditorfn", "refeditor0", "refcrossfn",
}

HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3",
    "Heading 4", "Heading 5", "Heading 6",
}

MIN_GROUP_SIZE = 2
# Increased from 10 to 15 to handle long zoning technical standards sections
# (e.g. Sec. 78-164 has 11 items, some sections have more)
MAX_ITEMS_IN_SUMMARY = 15


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _is_list_para(para) -> bool:
    style = para.style.name
    return style in LIST_STYLES and style not in SKIP_STYLES


def _is_heading(para) -> bool:
    style = para.style.name
    text = para.text.strip()
    if style in HEADING_STYLES:
        return True
    if style == "Normal" and re.match(r'^Sec\.\s+\d', text):
        return True
    return False


def _clean_item(text: str) -> str:
    """Strip leading list markers: (1)  a.  1.  •  -"""
    text = text.strip()
    text = re.sub(r'^\s*[\(\[]?[0-9a-zA-Z]{1,3}[\)\]\.][\s]+', '', text)
    text = re.sub(r'^\s*[•\-\*]\s+', '', text)
    return text.strip()


def _build_prose(items: list, heading: str) -> str:
    """Join list items into a single readable prose sentence."""
    cleaned = [_clean_item(it) for it in items if _clean_item(it)]
    if not cleaned:
        return ""

    capped = cleaned[:MAX_ITEMS_IN_SUMMARY]

    if len(capped) == 1:
        body = capped[0]
    elif len(capped) == 2:
        body = f"{capped[0]} and {capped[1]}"
    else:
        body = "; ".join(capped[:-1]) + "; and " + capped[-1]

    prefix = f"{heading}: " if heading else ""
    return f"{prefix}{body}."


def _insert_paragraph_after(doc: Document, para_index: int, text: str):
    """Insert a Normal-style paragraph after para_index using direct XML manipulation."""
    ref_para = doc.paragraphs[para_index]

    new_p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    new_p.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)

    ref_para._element.addnext(new_p)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PUBLIC FUNCTION
# ─────────────────────────────────────────────────────────────────────────────

def preprocess_document(filepath: str, doc_id: str = "") -> str:
    """
    Load a DOCX, inject prose summaries after all list groups,
    save to a temp file, and return the temp file path.

    The original source file is NEVER modified.

    ARGS:
        filepath:  Full path to the source DOCX (from Google Drive)
        doc_id:    e.g. "ch78" — used for log output only

    RETURNS:
        Path to temp DOCX with summaries injected.
        Pass this to chunk_document() in place of the original filepath.
        build_corpus.py deletes the temp file after chunking.
    """
    doc = Document(filepath)
    paras = doc.paragraphs
    n = len(paras)

    current_heading = ""
    injections = []   # (last_list_idx, prose_text)

    i = 0
    while i < n:
        para = paras[i]
        text = para.text.strip()
        style = para.style.name

        if not text or style in SKIP_STYLES:
            i += 1
            continue

        if _is_heading(para):
            current_heading = text
            i += 1
            continue

        if _is_list_para(para):
            items = []
            j = i
            while j < n and _is_list_para(paras[j]):
                t = paras[j].text.strip()
                if t:
                    items.append(t)
                j += 1

            last_idx = j - 1

            if len(items) >= MIN_GROUP_SIZE:
                prose = _build_prose(items, current_heading)
                if prose:
                    injections.append((last_idx, prose))

            i = j
            continue

        i += 1

    # Insert in reverse order to preserve paragraph indices
    for idx, prose in sorted(injections, key=lambda x: x[0], reverse=True):
        _insert_paragraph_after(doc, idx, prose)

    print(f"    [post_processor] {doc_id}: "
          f"injected {len(injections)} prose summaries")

    suffix = Path(filepath).suffix or ".docx"
    tmp = tempfile.NamedTemporaryFile(
        suffix=suffix,
        prefix=f"brentwood_{doc_id}_preprocessed_",
        delete=False,
    )
    tmp.close()
    doc.save(tmp.name)

    return tmp.name
