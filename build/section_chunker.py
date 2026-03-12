"""
section_chunker.py
==================
City of Brentwood Engineering AI Assistant - V2
Document chunking engine for the municipal corpus.

PURPOSE:
    Reads a DOCX file and breaks it into searchable text chunks.
    Every chunk carries metadata about exactly where it came from —
    which document, which article, which section — so the AI can
    cite sources precisely.

WHAT MAKES THIS CHUNKER SPECIAL:
    Most chunkers split text every N characters regardless of meaning.
    This chunker respects the legal structure of each document:
    - Chunks never cross section boundaries
    - Every chunk inherits its parent section's metadata
    - Tables are converted to prose so their data is searchable
    - History/ordinance notes are excluded (not policy content)

HANDLES THREE DOCUMENT TYPES:
    1. Municipal Code chapters (Ch. 56, Ch. 78, etc.)
       - Sections start with "Sec. 56-1." pattern
       - Organized into Articles and Divisions

    2. Appendix A - Subdivision Regulations
       - Sections start with "6.10 -" decimal pattern (no "Sec." prefix)
       - Organized into ARTICLE ONE, ARTICLE TWO, etc.

    3. Engineering Policy Manual
       - Internal policy sections identified by Heading 1 style
       - All-caps headings like "FOUNDATION SURVEYS", "RETAINING WALLS"
       - Contains embedded Municipal Code quotes (cited differently)

CHUNK METADATA FIELDS:
    Every chunk gets these fields stored in ChromaDB:
    - doc_id:           short ID from registry ("ch56", "appendix_a", "epm")
    - doc_title:        full human-readable title
    - content_type:     "municipal_code" | "appendix" | "engineering_policy" |
                        "code_reference" | "external_reference"
    - article:          e.g. "ARTICLE I" or "ARTICLE ONE" or "" if none
    - division:         e.g. "DIVISION 2" or "" if none
    - section_number:   e.g. "56-31" or "6.10" or "FOUNDATION SURVEYS"
    - section_title:    e.g. "Water quality riparian buffer"
    - chunk_index:      position of this chunk within the section (0, 1, 2...)
    - chunk_total:      total chunks in this section
    - source_citation:  pre-formatted citation string for display

AUTHOR:  City of Brentwood Engineering Department AI Assistant Project
VERSION: 2.0
"""

import re
from dataclasses import dataclass, field
from typing import Optional
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# DATA STRUCTURES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class DocumentSection:
    """
    Represents one logical section of a document before chunking.

    Think of this as a folder that holds all the text belonging to
    one section (e.g., Sec. 56-31 Water quality riparian buffer).
    After collection, it gets split into one or more chunks.
    """
    # Location metadata — where does this section live in the document?
    article: str = ""           # e.g. "ARTICLE I" or "ARTICLE ONE"
    division: str = ""          # e.g. "DIVISION 2 — EROSION AND SEDIMENT CONTROL"
    section_number: str = ""    # e.g. "56-31" or "6.10" or "FOUNDATION SURVEYS"
    section_title: str = ""     # e.g. "Water quality riparian buffer"
    content_type: str = ""      # see module docstring for valid values

    # The actual text content of this section (all paragraphs joined)
    text: str = ""


@dataclass
class Chunk:
    """
    A single searchable unit of text ready for insertion into ChromaDB.

    This is the final output of the chunker. Each chunk gets:
    - A unique ID used as the ChromaDB record key
    - The text that gets embedded into a vector
    - Metadata that travels with the chunk forever
    """
    chunk_id: str               # Unique identifier: "ch56_sec56-31_0"
    text: str                   # The actual text to embed and search
    metadata: dict              # All the fields listed in module docstring


# ─────────────────────────────────────────────────────────────────────────────
# SECTION DETECTION PATTERNS
# ─────────────────────────────────────────────────────────────────────────────

# Municipal Code section header: "Sec. 56-31. - Water quality riparian buffer."
# Captures: group(1) = "56-31", group(2) = "Water quality riparian buffer"
MUNICODE_SEC_PATTERN = re.compile(
    r'^Sec\.\s+([\d]+[-\w\.]+)\.\s*[-—]?\s*(.*)',
    re.IGNORECASE
)

# Appendix A section header: "6.10 - Storm Drainage." or "3.5 - Preliminary Plan Contents."
# Captures: group(1) = "6.10", group(2) = "Storm Drainage"
APPENDIX_SEC_PATTERN = re.compile(
    r'^(\d+\.\d+)\s*[-—]\s*(.*)'
)

# Municipal Code article: "ARTICLE I. - STORMWATER MANAGEMENT AND EROSION CONTROL"
MUNICODE_ARTICLE_PATTERN = re.compile(
    r'^ARTICLE\s+([\w]+)\.\s*[-—]?\s*(.*)',
    re.IGNORECASE
)

# Appendix A article: "ARTICLE ONE. - GENERAL PROVISIONS"
APPENDIX_ARTICLE_PATTERN = re.compile(
    r'^ARTICLE\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\.\s*[-—]?\s*(.*)',
    re.IGNORECASE
)

# Municipal Code division: "DIVISION 2. - EROSION AND SEDIMENT CONTROL"
DIVISION_PATTERN = re.compile(
    r'^DIVISION\s+([\w]+)\.\s*[-—]?\s*(.*)',
    re.IGNORECASE
)

# "From 78-45:" or "From EC Checklist" or "From TDEC:" — EPM source attribution
FROM_PATTERN = re.compile(
    r'^From\s+(.+?)[:,]?\s*$',
    re.IGNORECASE
)

# History/ordinance note — these are amendment records, not policy content
# "Ord. No. 2008-02, § 1, 3-24-2008" or "(Amend. of 8-5-2024)"
HISTORY_NOTE_PATTERN = re.compile(
    r'^(\(Ord\.|Ord\.\s+No\.|Amend\.\s+of)',
    re.IGNORECASE
)

# Styles to exclude entirely from chunking
EXCLUDED_STYLES = {
    'historynote0',   # Ordinance amendment records in Municipal Code
    'refeditorfn',    # Editor's footnotes in Municipal Code
    'refeditor0',     # Editor reference notes
}

# CHUNK SIZE SETTINGS
# These values balance two competing needs:
# - Chunks too small: miss context, poor retrieval accuracy
# - Chunks too large: dilute relevance score, slow embedding
CHUNK_TARGET_CHARS = 600       # aim for chunks around this size
CHUNK_MAX_CHARS = 900          # hard ceiling before forced split
CHUNK_OVERLAP_CHARS = 120      # overlap between consecutive chunks in same section


# ─────────────────────────────────────────────────────────────────────────────
# TABLE CONVERSION — CRITICAL FOR ENGINEERING DATA
# ─────────────────────────────────────────────────────────────────────────────

def table_to_prose(table, table_name: str = "", section_context: str = "") -> str:
    """
    Convert a DOCX table into a searchable prose description.

    WHY THIS EXISTS:
        python-docx's paragraph iterator skips table content entirely.
        If we don't explicitly extract tables, engineering values like
        "minimum easement width for 36-inch pipe at 8-foot depth = 25 feet"
        simply do not exist in our vector database. Engineers querying
        those values would get "not found" or a hallucinated answer.

    HOW IT WORKS:
        Reads every cell, identifies header rows vs data rows, and handles
        two common table layouts:

        Layout A — Row-label table (first column header is empty):
            | (empty)       | Avg Width | Min Width | Notes |
            | Standard waters | 30       | 15        | ...   |
            | Exceptional ETW | 60       | 30        | ...   |
        → Output: "Standard waters; Average Buffer Width = 30; Minimum Buffer Width = 15"

        Layout B — Standard two-column lookup table:
            | Channel Width     | Required Easement Width  |
            | Less than 5 feet  | 10 feet                  |
        → Output: "Channel Width = Less than 5 feet; Required Easement Width = 10 feet"

    ARGS:
        table:           python-docx Table object
        table_name:      e.g. "TABLE SIX" (extracted from surrounding text)
        section_context: e.g. "Section 6.10 — Storm Drainage" for context

    RETURNS:
        A prose string containing all table data in natural language.
    """
    if not table.rows:
        return ""

    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(c for c in cells):
            rows.append(cells)

    if not rows:
        return ""

    lines = []

    # Header line
    header_parts = []
    if table_name:
        header_parts.append(table_name)
    if section_context:
        header_parts.append(f"({section_context})")
    if header_parts:
        lines.append(" ".join(header_parts) + ":")

    header_row = rows[0] if rows else []
    data_rows = rows[1:] if len(rows) > 1 else []

    if not data_rows:
        # Single-row table — output as a plain statement
        if header_row:
            lines.append(". ".join(c for c in header_row if c) + ".")
        return "\n".join(lines)

    # Detect row-label layout: first column header is empty, meaning col 0
    # contains the row subject label and cols 1+ contain measured values.
    # Example: riparian buffer table where col[0]="" and col[1]="Avg Width".
    has_row_label = bool(header_row) and (not header_row[0])
    col_headers = header_row[1:] if has_row_label else header_row

    for row_cells in data_rows:
        parts = []

        if has_row_label:
            # First cell is the subject of this row (e.g. "Standard waters")
            row_label = row_cells[0] if row_cells else ""
            values = row_cells[1:]
            if row_label:
                parts.append(row_label)
        else:
            values = row_cells

        # Pair each value with its column header
        for i, val in enumerate(values):
            if not val:
                continue
            if i < len(col_headers) and col_headers[i]:
                # Don't repeat the header as the value (merged cell artifact)
                if val.lower() != col_headers[i].lower():
                    parts.append(f"{col_headers[i]} = {val}")
            else:
                # No column header for this position — output value alone
                parts.append(val)

        if parts:
            lines.append("; ".join(parts) + ".")

    return "\n".join(lines)


def extract_table_name_from_text(text: str) -> str:
    """
    Look for a table name label in text like "TABLE SIX" or "TABLE 6".
    Returns the matched name, or empty string if not found.
    """
    match = re.search(
        r'\bTABLE\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|ELEVEN|\d+)\b',
        text,
        re.IGNORECASE
    )
    return match.group(0).upper() if match else ""


# ─────────────────────────────────────────────────────────────────────────────
# MAIN CHUNKER CLASS
# ─────────────────────────────────────────────────────────────────────────────

class SectionChunker:
    """
    Converts a DOCX file into a list of Chunk objects ready for ChromaDB.

    USAGE:
        chunker = SectionChunker()
        chunks = chunker.chunk_document(
            filepath="/path/to/Chapter_56.docx",
            doc_id="ch56",
            doc_title="Brentwood Municipal Code Chapter 56 — Stormwater Management",
            doc_type="municipal_code",
            citation_format="Brentwood Municipal Code, Chapter 56"
        )

    The doc_type parameter controls which section detection pattern is used:
        "municipal_code"      — Sec. XX-XX pattern
        "appendix"            — X.X pattern (Appendix A)
        "engineering_policy"  — Heading 1 style sections
    """

    # Paragraphs whose ALL-CAPS text is document-level boilerplate,
    # not a policy section heading. Excluded from EPM section detection.
    _EPM_TITLE_EXCLUSIONS = re.compile(
        r'^(CITY OF|ENGINEERING DEPARTMENT|ENGINEERING REVIEW|'
        r'REFER ALSO TO|FURTHERMORE)',
        re.IGNORECASE
    )
    # Timestamp pattern — e.g. "(11/23/2015 9:12 AM MAW)"
    _TIMESTAMP_PATTERN = re.compile(r'^\(\d{1,2}/\d{1,2}/\d{4}')

    def __init__(self):
        # Running state while processing a document
        self._current_article = ""
        self._current_division = ""
        self._doc_id = ""
        self._doc_title = ""
        self._doc_type = ""
        self._citation_format = ""
        self._pending_table_name = ""  # table name seen in text before the table
        # Tracks how many times each section_number has appeared in this doc.
        # The EPM quotes the same Municipal Code section multiple times.
        # Without this counter, chunk IDs collide and ChromaDB silently
        # overwrites earlier chunks with later ones.
        self._section_occurrence_counter = {}

    def chunk_document(
        self,
        filepath: str,
        doc_id: str,
        doc_title: str,
        doc_type: str,
        citation_format: str
    ) -> list:
        """
        Main entry point. Reads a DOCX and returns a list of Chunk objects.

        ARGS:
            filepath:        full path to the DOCX file
            doc_id:          short identifier from registry (e.g. "ch56")
            doc_title:       full title for display (e.g. "Chapter 56 — Stormwater")
            doc_type:        "municipal_code" | "appendix" | "engineering_policy"
            citation_format: base citation string (e.g. "Brentwood Municipal Code, Chapter 56")

        RETURNS:
            List of Chunk objects, ready for ChromaDB insertion.
        """
        # Store document-level state
        self._doc_id = doc_id
        self._doc_title = doc_title
        self._doc_type = doc_type
        self._citation_format = citation_format
        self._current_article = ""
        self._current_division = ""
        self._pending_table_name = ""
        self._section_occurrence_counter = {}  # reset for each new document

        # Open the document
        doc = Document(filepath)

        # Step 1: Collect all sections (text grouped by section boundaries)
        sections = self._collect_sections(doc)

        # Step 2: Split each section into chunks with overlap
        all_chunks = []
        for section in sections:
            section_chunks = self._split_section_into_chunks(section)
            all_chunks.extend(section_chunks)

        return all_chunks

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION COLLECTION
    # ─────────────────────────────────────────────────────────────────────────

    def _collect_sections(self, doc) -> list:
        """
        Walk the document paragraph by paragraph, grouping text into sections.

        This is like a filing clerk reading a legal document and putting
        each section's pages into its own labeled folder.

        For tables: converts each table to prose and inserts it into the
        current section's text immediately after the surrounding paragraphs.
        """
        sections = []
        current_section = None

        # We need to interleave paragraphs and tables in document order.
        # python-docx's doc.paragraphs skips tables; doc.tables skips paragraphs.
        # The only way to get true document order is to iterate doc.element.body children.
        from docx.oxml.ns import qn

        body_children = list(doc.element.body)

        for child in body_children:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                # It's a paragraph — process it normally
                # Reconstruct a paragraph-like object from the XML element
                para_text = "".join(
                    node.text or ""
                    for node in child.iter()
                    if node.tag.endswith('}t')
                )
                para_text = para_text.strip()

                # Get the style name
                style_name = self._get_style_name(child)

                if not para_text:
                    continue  # skip blank paragraphs

                # Skip excluded styles (history notes, editor footnotes)
                if style_name in EXCLUDED_STYLES:
                    continue

                # Skip ordinance history notes by content pattern
                if HISTORY_NOTE_PATTERN.match(para_text):
                    continue

                # Check if this paragraph is a TABLE reference label
                # e.g. "TABLE SIX" appearing just before a table element
                table_name = extract_table_name_from_text(para_text)
                if table_name:
                    self._pending_table_name = table_name

                # Check for structural markers based on document type
                section_result = self._try_start_new_section(
                    para_text, style_name, current_section
                )

                if section_result is not None:
                    # This paragraph starts a new section
                    if current_section is not None and current_section.text.strip():
                        sections.append(current_section)
                    current_section = section_result
                else:
                    # This paragraph belongs to the current section
                    if current_section is None:
                        # Preamble text before any section — create a catch-all
                        current_section = DocumentSection(
                            section_number="preamble",
                            section_title="Document Preamble",
                            content_type=self._base_content_type(),
                        )
                    current_section.text += para_text + " "

            elif tag == 'tbl':
                # It's a table — convert to prose and append to current section
                if current_section is None:
                    current_section = DocumentSection(
                        section_number="preamble",
                        section_title="Document Preamble",
                        content_type=self._base_content_type(),
                    )

                table_prose = self._convert_table_element(
                    child,
                    doc,
                    current_section
                )
                if table_prose:
                    current_section.text += "\n" + table_prose + "\n"
                self._pending_table_name = ""  # consumed

        # Don't forget the last section
        if current_section is not None and current_section.text.strip():
            sections.append(current_section)

        return sections

    def _get_style_name(self, para_element) -> str:
        """Extract the style name from a paragraph XML element."""
        from docx.oxml.ns import qn
        pPr = para_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is None:
            return 'Normal'
        pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        if pStyle is None:
            return 'Normal'
        return pStyle.get(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
            'Normal'
        )

    def _convert_table_element(self, tbl_element, doc, current_section) -> str:
        """
        Convert a table XML element to a prose string.

        Reconstructs a python-docx Table object from the element so we can
        use the same table_to_prose() function.
        """
        from docx.table import Table
        from docx.oxml import OxmlElement

        try:
            table = Table(tbl_element, doc)
            section_context = ""
            if current_section:
                section_context = f"Section {current_section.section_number}"
                if current_section.section_title:
                    section_context += f" — {current_section.section_title}"

            return table_to_prose(
                table,
                table_name=self._pending_table_name,
                section_context=section_context
            )
        except Exception as e:
            # Never fail silently on table conversion — log and continue
            return f"[Table data — conversion error: {str(e)[:80]}]"

    def _try_start_new_section(
        self,
        text: str,
        style_name: str,
        current_section
    ) -> Optional[DocumentSection]:
        """
        Determine if this paragraph starts a new section.

        Returns a new DocumentSection if yes, None if this paragraph
        should be added to the current section.

        Also updates self._current_article and self._current_division
        as structural markers are encountered.
        """
        # ── Check for article-level markers (update state, don't create section) ──

        art_match = MUNICODE_ARTICLE_PATTERN.match(text)
        if not art_match:
            art_match = APPENDIX_ARTICLE_PATTERN.match(text)

        if art_match:
            # Update the running article tracker
            # Clean up the article text — remove footnote markers like [6]
            article_text = re.sub(r'\[\d+\]', '', text).strip()
            self._current_article = article_text
            self._current_division = ""  # new article resets division
            return None  # article headers don't become their own sections

        div_match = DIVISION_PATTERN.match(text)
        if div_match:
            division_text = re.sub(r'\[\d+\]', '', text).strip()
            self._current_division = division_text
            return None  # division headers don't become their own sections

        # ── Check for section-level markers (create new section) ──

        if self._doc_type == "municipal_code":
            return self._detect_municode_section(text, style_name)

        elif self._doc_type == "appendix":
            return self._detect_appendix_section(text, style_name)

        elif self._doc_type == "engineering_policy":
            return self._detect_epm_section(text, style_name)

        return None

    def _detect_municode_section(self, text: str, style: str) -> Optional[DocumentSection]:
        """Detect section starts in Municipal Code chapters."""
        match = MUNICODE_SEC_PATTERN.match(text)
        if match and style == 'Normal':
            sec_num_raw = match.group(1)
            # Clean title: remove trailing period and footnote markers
            sec_title = re.sub(r'\[\d+\]', '', match.group(2)).strip().rstrip('.')
            # Route through uniqueness check — handles duplicate section headers
            # that can appear when post_processor injects prose paragraphs.
            sec_num = self._make_unique_section_number(sec_num_raw)
            return DocumentSection(
                article=self._current_article,
                division=self._current_division,
                section_number=sec_num,
                section_title=sec_title,
                content_type="municipal_code",
                text=""  # text gets filled in as paragraphs are added
            )
        return None

    def _detect_appendix_section(self, text: str, style: str) -> Optional[DocumentSection]:
        """Detect section starts in Appendix A."""
        match = APPENDIX_SEC_PATTERN.match(text)
        if match:
            sec_num_raw = match.group(1)
            sec_title = re.sub(r'\[\d+\]', '', match.group(2)).strip().rstrip('.')
            # Route through uniqueness check — same reason as municode above.
            sec_num = self._make_unique_section_number(sec_num_raw)
            return DocumentSection(
                article=self._current_article,
                division=self._current_division,
                section_number=sec_num,
                section_title=sec_title,
                content_type="appendix",
                text=""
            )
        return None

    # Pattern to extract section number from EPM headings like:
    #   "9. RESIDENTIAL DRIVEWAYS (Sec. 78-486)"  →  num="9", title="RESIDENTIAL DRIVEWAYS"
    #   "10.1 Height Limits"                       →  num="10.1", title="Height Limits"
    #   "2.2.1 Site Plan Requirements"             →  num="2.2.1", title="Site Plan Requirements"
    _EPM_HEADING_PATTERN = re.compile(r'^(\d+(?:\.\d+)*\.?)\s+(.+)$')

    def _detect_epm_section(self, text: str, style: str) -> Optional[DocumentSection]:
        """
        Detect section starts in the Engineering Policy Manual.

        The EPM (Version 3.0, March 2026) uses a clean numbered heading hierarchy:
          Heading 1  — Major sections:    "10. RETAINING WALLS"
          Heading 2  — Subsections:       "10.1 Height Limits"
          Heading 3  — Sub-subsections:   "2.2.1 Site Plan Requirements"

        Each Heading 1/2/3 paragraph starts a new chunk with content_type
        "engineering_policy".

        The EPM also embeds quoted Municipal Code text (Normal style, starts
        with "Sec. XX-XX") and "From" attribution blocks — these are kept as
        separate chunk types for citation purposes.

        NOTE: The old EPM used ALL-CAPS Normal paragraphs as headings.
        Version 3.0 no longer uses that pattern — heading detection is now
        purely style-based (Heading 1/2/3). The ALL-CAPS Normal fallback is
        intentionally removed to prevent formula lines like
        "R = A + B - (A × B) / 100" from being misdetected as headings.
        """
        # ── Heading 1 / 2 / 3 — numbered EPM policy sections ─────────────
        if style in ('Heading 1', 'Heading 2', 'Heading 3',
                     'Heading1', 'Heading2', 'Heading3'):
            heading_match = self._EPM_HEADING_PATTERN.match(text.strip())
            if heading_match:
                # e.g. "10. RETAINING WALLS (Sec. 78-14)" →
                #   sec_num = "10", sec_title = "RETAINING WALLS"
                raw_num   = heading_match.group(1).rstrip('.')
                raw_title = heading_match.group(2).strip()

                # Strip trailing code references like "(Sec. 78-14)" or
                # "(Subdivision Regulations 6.10)" from the title — these
                # are navigation aids, not part of the section name
                sec_title = re.sub(
                    r'\s*\([^)]*(?:Sec\.|Subdivision|Regulations|Division)[^)]*\)',
                    '', raw_title
                ).strip()

                # Use the dotted number as the section identifier so that
                # "10.1" is clearly a child of "10" in citations
                sec_num = self._make_unique_section_number(raw_num)

                # Track article (Heading 1) and division (Heading 2) for
                # context inheritance by child chunks
                if style in ('Heading 1', 'Heading1'):
                    self._current_article   = sec_title
                    self._current_division  = ""
                elif style in ('Heading 2', 'Heading2'):
                    self._current_division  = sec_title

                return DocumentSection(
                    article=self._current_article,
                    division=self._current_division,
                    section_number=sec_num,
                    section_title=sec_title,
                    content_type="engineering_policy",
                    text=""
                )
            else:
                # Heading style but no number prefix (e.g. "APPENDIX A: CODE REFERENCE SUMMARY")
                # Treat as a top-level policy section using the full text as title
                sec_title = text.strip().rstrip(':')
                sec_num   = self._make_unique_section_number(
                    re.sub(r'[^a-z0-9]', '_', sec_title.lower())[:40]
                )
                if style in ('Heading 1', 'Heading1'):
                    self._current_article  = sec_title
                    self._current_division = ""
                return DocumentSection(
                    article=self._current_article,
                    division=self._current_division,
                    section_number=sec_num,
                    section_title=sec_title,
                    content_type="engineering_policy",
                    text=""
                )

        # ── Quoted Municipal Code sections inside the EPM ─────────────────
        # These appear as Normal-style paragraphs starting with "Sec. XX-XX"
        municode_match = MUNICODE_SEC_PATTERN.match(text)
        if municode_match:
            sec_num_raw = municode_match.group(1)
            sec_title = re.sub(r'\[\d+\]', '', municode_match.group(2)).strip().rstrip('.')
            sec_num = self._make_unique_section_number(sec_num_raw)
            return DocumentSection(
                article=self._current_article,
                division="",
                section_number=sec_num,
                section_title=sec_title,
                content_type="code_reference",
                text=text + " "  # include the header line as text too
            )

        # ── "From" attribution blocks ─────────────────────────────────────
        from_match = FROM_PATTERN.match(text)
        if from_match:
            source = from_match.group(1).strip()
            sec_num_raw = f"from_{re.sub(r'[^a-z0-9]', '_', source.lower())}"
            sec_num = self._make_unique_section_number(sec_num_raw)
            return DocumentSection(
                article="",
                division="",
                section_number=sec_num,
                section_title=f"Reference: {source}",
                content_type="external_reference",
                text=text + " "
            )

        return None

    def _make_unique_section_number(self, sec_num: str) -> str:
        """
        Ensure section numbers are unique within a document.

        The EPM quotes the same Municipal Code section multiple times
        (e.g. Sec. 58-6 appears 3 times, Sec. 78-486 appears 5 times).
        If we don't deduplicate these, all occurrences get the same
        chunk_id and ChromaDB silently overwrites earlier ones.

        First occurrence:  "78-486"       (no suffix, clean citation)
        Second occurrence: "78-486_occ2"
        Third occurrence:  "78-486_occ3"
        """
        key = sec_num
        self._section_occurrence_counter[key] = \
            self._section_occurrence_counter.get(key, 0) + 1
        count = self._section_occurrence_counter[key]
        if count == 1:
            return sec_num
        else:
            return f"{sec_num}_occ{count}"

    def _base_content_type(self) -> str:
        """Default content type for preamble/uncategorized content."""
        type_map = {
            "municipal_code": "municipal_code",
            "appendix": "appendix",
            "engineering_policy": "engineering_policy",
        }
        return type_map.get(self._doc_type, "municipal_code")

    # ─────────────────────────────────────────────────────────────────────────
    # CHUNK SPLITTING
    # ─────────────────────────────────────────────────────────────────────────

    def _split_section_into_chunks(self, section: DocumentSection) -> list:
        """
        Split a section's text into overlapping chunks.

        STRATEGY:
            - If section text is short (< CHUNK_MAX_CHARS): single chunk
            - If section text is long: split at sentence boundaries,
              aiming for CHUNK_TARGET_CHARS per chunk, with
              CHUNK_OVERLAP_CHARS of text repeated at each boundary

        WHY OVERLAP:
            If an answer spans the split point between two chunks,
            overlap ensures that at least one chunk contains the full
            relevant sentence.

        RETURNS:
            List of Chunk objects, each with a unique ID and metadata.
        """
        text = section.text.strip()
        if not text:
            return []

        # Split text into sentence-ish units for clean boundaries
        # Split on ". " or ".\n" but keep the period with the preceding sentence
        sentences = self._split_into_sentences(text)

        if not sentences:
            return []

        # If the whole section fits in one chunk, no need to split
        if len(text) <= CHUNK_MAX_CHARS:
            return [self._make_chunk(section, text, 0, 1)]

        # Build chunks by accumulating sentences until size limit is reached
        chunks_text = []
        current_chunk = []
        current_len = 0

        for sentence in sentences:
            sentence_len = len(sentence)

            if current_len + sentence_len > CHUNK_MAX_CHARS and current_chunk:
                # Current chunk is full — save it and start a new one
                chunk_text = " ".join(current_chunk).strip()
                chunks_text.append(chunk_text)

                # Start new chunk with overlap from end of current chunk
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + [sentence]
                current_len = sum(len(s) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_len += sentence_len

        # Don't forget the last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk).strip()
            chunks_text.append(chunk_text)

        # Deduplicate — sometimes overlap creates near-identical consecutive chunks
        chunks_text = self._deduplicate_chunks(chunks_text)

        # Convert text list to Chunk objects with metadata
        total = len(chunks_text)
        return [
            self._make_chunk(section, text, i, total)
            for i, text in enumerate(chunks_text)
        ]

    def _split_into_sentences(self, text: str) -> list:
        """
        Split text into sentence-level units for clean chunk boundaries.

        We split on sentence-ending punctuation followed by whitespace,
        keeping the punctuation with the sentence it ends.
        """
        # Split on ". " or ".\n" keeping the period
        # Also split on subsection labels like "(a) " at the start
        parts = re.split(r'(?<=[.!?])\s+', text)
        # Filter empty strings
        return [p.strip() for p in parts if p.strip()]

    def _get_overlap_text(self, sentences: list) -> list:
        """
        Get the last N sentences whose total length is ~CHUNK_OVERLAP_CHARS.
        These will be prepended to the next chunk for context continuity.
        """
        overlap = []
        total_len = 0
        for sentence in reversed(sentences):
            if total_len + len(sentence) > CHUNK_OVERLAP_CHARS:
                break
            overlap.insert(0, sentence)
            total_len += len(sentence)
        return overlap

    def _deduplicate_chunks(self, chunks: list) -> list:
        """Remove consecutive chunks that are too similar (>90% identical)."""
        if len(chunks) <= 1:
            return chunks

        result = [chunks[0]]
        for chunk in chunks[1:]:
            prev = result[-1]
            # Check overlap: if new chunk starts with most of the previous chunk's end
            overlap_check = prev[-100:] if len(prev) > 100 else prev
            if overlap_check not in chunk:
                result.append(chunk)
            # If they overlap heavily, skip the duplicate
        return result

    def _make_chunk(
        self,
        section: DocumentSection,
        text: str,
        chunk_index: int,
        chunk_total: int
    ) -> Chunk:
        """
        Construct a Chunk object with full metadata for ChromaDB insertion.

        The chunk_id must be unique across the entire corpus.
        Format: "{doc_id}_sec{section_number}_{chunk_index}"
        Example: "ch56_sec56-31_0", "appendix_a_sec6.10_2", "epm_secFOUNDATION_SURVEYS_0"
        """
        # Build safe section ID for use in chunk_id
        safe_sec = re.sub(r'[^a-z0-9]', '_', section.section_number.lower())
        chunk_id = f"{self._doc_id}_sec{safe_sec}_{chunk_index}"

        # Build the citation string for this specific section
        source_citation = self._build_citation(section)

        metadata = {
            "doc_id": self._doc_id,
            "doc_title": self._doc_title,
            "content_type": section.content_type,
            "article": section.article,
            "division": section.division,
            "section_number": section.section_number,
            "section_title": section.section_title,
            "chunk_index": chunk_index,
            "chunk_total": chunk_total,
            "source_citation": source_citation,
        }

        return Chunk(
            chunk_id=chunk_id,
            text=text,
            metadata=metadata
        )

    def _build_citation(self, section: DocumentSection) -> str:
        """
        Build a human-readable citation string for a section.

        CITATION FORMATS BY CONTENT TYPE:

        municipal_code:
            "Brentwood Municipal Code, Chapter 56 — Stormwater Management,
             Sec. 56-31 — Water quality riparian buffer"

        appendix:
            "City of Brentwood Subdivision Regulations (Appendix A),
             Section 6.10 — Storm Drainage"

        engineering_policy:
            "City of Brentwood Engineering Dept. Policy Manual —
             FOUNDATION SURVEYS"

        code_reference:
            "Brentwood Municipal Code, Sec. 78-45 (referenced in
             Engineering Policy Manual)"

        external_reference:
            "Reference: EC Checklist (cited in Engineering Policy Manual)"
        """
        ct = section.content_type
        sec_num = section.section_number
        sec_title = section.section_title

        if ct == "municipal_code":
            base = self._citation_format  # e.g. "Brentwood Municipal Code, Chapter 56 — Stormwater"
            if sec_num and sec_title:
                return f"{base}, Sec. {sec_num} — {sec_title}"
            elif sec_num:
                return f"{base}, Sec. {sec_num}"
            else:
                return base

        elif ct == "appendix":
            base = self._citation_format  # e.g. "City of Brentwood Subdivision Regulations (Appendix A)"
            if sec_num and sec_title:
                return f"{base}, Section {sec_num} — {sec_title}"
            elif sec_num:
                return f"{base}, Section {sec_num}"
            else:
                return base

        elif ct == "engineering_policy":
            # Include the dotted section number for precision:
            #   "City of Brentwood Engineering Dept. Policy Manual — Sec. 10.1: Height Limits"
            # For top-level sections without a numeric prefix (e.g. appendix headers),
            # fall back to title only.
            if sec_num and re.match(r'^\d', sec_num):
                return (
                    f"City of Brentwood Engineering Dept. Policy Manual — "
                    f"Sec. {sec_num}: {sec_title}"
                )
            else:
                return f"City of Brentwood Engineering Dept. Policy Manual — {sec_title}"

        elif ct == "code_reference":
            return (
                f"Brentwood Municipal Code, Sec. {sec_num} — {sec_title} "
                f"(referenced in Engineering Policy Manual)"
            )

        elif ct == "external_reference":
            return f"{sec_title} (cited in Engineering Policy Manual)"

        else:
            return self._citation_format


# ─────────────────────────────────────────────────────────────────────────────
# CONVENIENCE FUNCTION
# ─────────────────────────────────────────────────────────────────────────────

def chunk_document(
    filepath: str,
    doc_id: str,
    doc_title: str,
    doc_type: str,
    citation_format: str
) -> list:
    """
    Module-level convenience function.
    Equivalent to SectionChunker().chunk_document(...).

    This is what build_corpus.py calls for each document:

        from section_chunker import chunk_document
        chunks = chunk_document(
            filepath=str(doc_path),
            doc_id=entry["doc_id"],
            doc_title=entry["title"],
            doc_type=entry["doc_type"],
            citation_format=entry["citation_format"]
        )
    """
    chunker = SectionChunker()
    return chunker.chunk_document(
        filepath=filepath,
        doc_id=doc_id,
        doc_title=doc_title,
        doc_type=doc_type,
        citation_format=citation_format
    )
