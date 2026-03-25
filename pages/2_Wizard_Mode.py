"""
2_Wizard_Mode.py
================
City of Brentwood Engineering AI Assistant - V2
Wizard Mode — interactive plan review checklist with export.
"""

import streamlit as st
import sys
import re
import csv
import zlib
import random
import string
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

# Brentwood, TN — Central Time (handles CST/CDT automatically)
_CT = ZoneInfo("America/Chicago")
from io import BytesIO, StringIO

# pages/ is one level down — add both utils/ and repo root
sys.path.append(str(Path(__file__).parent.parent / "utils"))
sys.path.append(str(Path(__file__).parent.parent))

from checklist_data import (
    REVIEW_TYPES,
    REVIEWERS,
    CHECKLIST_SECTIONS,
    get_checklist_for_review_type
)
from comments_database import COMMENTS, get_comment
from theme import apply_theme, render_sidebar, page_header, section_heading, footer, get_favicon

# ── Traffic Calming module (gracefully unavailable if files not yet uploaded) ──
try:
    from traffic_calming_data import (
        ARTERIAL_STREETS,
        COLLECTOR_STREETS,
        STREET_CLASSIFICATIONS,
        APPLICATION_TYPES,
        TIER2_STRATEGIES,
        SCORING_CRITERIA,
    )
    from traffic_calming_report import build_traffic_calming_report
    TC_AVAILABLE = True
except ImportError:
    TC_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="Wizard Mode — Brentwood Engineering AI",
                   page_icon=get_favicon(), layout="wide")

apply_theme()
render_sidebar(active="wizard")

# Wizard-specific CSS: comment boxes, back-to-top button, button active states
st.markdown("""
<style>
    /* ── Step headings (Step 1, Step 2, Step 3) ──────────────────────── */
    .bw-step-heading {
        font-family: var(--font-sans);
        font-size: 1.35rem;
        font-weight: 700;
        color: #22427C;
        border-left: 4px solid #F07138;
        padding: 0.4rem 0 0.4rem 0.75rem;
        margin: 1.5rem 0 1rem 0;
        line-height: 1.3;
    }
    /* ── Checklist section header ─────────────────────────────────────── */
    .bw-section-header {
        background: #EEF2F9 !important;
        color: #22427C !important;
        padding: 0.65rem 1rem;
        border-left: 4px solid #F07138;
        margin: 1.2rem 0 0.5rem 0;
        font-size: 0.97rem;
        font-weight: 700;
        letter-spacing: 0.02em;
        border-radius: 0 6px 6px 0;
    }
    /* ── Comment box (shown when No selected) ────────────────────────── */
    .bw-comment-box {
        background: #FFFBF0 !important;
        color: #1A2332 !important;
        border: 1px solid #F6D860;
        border-left: 3px solid #E8A000;
        border-radius: 6px;
        padding: 0.75rem 1rem;
        margin: 0.4rem 0 0.4rem 1.5rem;
        font-size: 0.9em;
    }
    .bw-resubmittal-box {
        background: #EEF2F9 !important;
        color: #1A2332 !important;
        border: 2px solid #2F5C9C;
        border-radius: 8px;
        padding: 1rem 1.2rem;
        margin: 1rem 0;
    }
    .bw-export-section {
        background: #F0FDF4 !important;
        color: #1A2332 !important;
        border: 1px solid #BBF7D0;
        border-left: 4px solid #16A34A;
        border-radius: 8px;
        padding: 1rem 1.2rem;
        margin: 1rem 0;
    }
    /* ── Back-to-top floating button ─────────────────────────────────── */
    #btt-btn {
        position: fixed;
        bottom: 2.5rem;
        right: 2rem;
        z-index: 9999;
        background: #22427C;
        color: white !important;
        border: none;
        border-radius: 50%;
        width: 44px;
        height: 44px;
        font-size: 1.3rem;
        line-height: 44px;
        text-align: center;
        cursor: pointer;
        box-shadow: 0 2px 8px rgba(34,66,124,0.35);
        opacity: 0;
        transition: opacity 0.25s;
        text-decoration: none;
    }
    #btt-btn.visible { opacity: 1; }
    #btt-btn:hover { background: #2F5C9C !important; }
</style>
<!-- Back-to-top anchor + button -->
<div id="page-top"></div>
<a id="btt-btn" href="#page-top" title="Back to top">↑</a>
<script>
    window.addEventListener('scroll', function() {
        var btn = document.getElementById('btt-btn');
        if (btn) {
            btn.classList.toggle('visible', window.scrollY > 300);
        }
    }, {passive: true});
</script>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'wizard_review_type' not in st.session_state:
        st.session_state.wizard_review_type = None
    if 'wizard_permit_number' not in st.session_state:
        st.session_state.wizard_permit_number = ""
    if 'wizard_address' not in st.session_state:
        st.session_state.wizard_address = ""
    if 'wizard_reviewer' not in st.session_state:
        st.session_state.wizard_reviewer = None
    if 'wizard_checklist_state' not in st.session_state:
        st.session_state.wizard_checklist_state = {}
    if 'wizard_selected_comments' not in st.session_state:
        st.session_state.wizard_selected_comments = {}
    if 'wizard_custom_notes' not in st.session_state:
        st.session_state.wizard_custom_notes = {}
    if 'wizard_started' not in st.session_state:
        st.session_state.wizard_started = False
    if 'wizard_resubmittal' not in st.session_state:
        st.session_state.wizard_resubmittal = "—"


def reset_checklist():
    """Reset checklist state when review type changes"""
    st.session_state.wizard_checklist_state = {}
    st.session_state.wizard_selected_comments = {}
    st.session_state.wizard_custom_notes = {}
    st.session_state.wizard_resubmittal = "—"


def collect_all_comments():
    """
    Collect all comments from checklist items marked 'No', then append
    the resubmittal comment (BB-0045) at the very end if the standalone
    resubmittal question was answered 'Yes'.

    Returns a list of raw comment strings (no prefix codes).
    Used by all export functions to ensure consistent output.
    """
    checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
    all_comments = []

    # Gather comments from all checklist items marked "No"
    for section_id, section_data in checklist.items():
        for item in section_data["items"]:
            item_key = item["id"]
            if st.session_state.wizard_checklist_state.get(item_key) == "No":
                selected = st.session_state.wizard_selected_comments.get(item_key, [])
                custom_note = st.session_state.wizard_custom_notes.get(item_key, "")

                for comment_id in selected:
                    comment_text = COMMENTS.get(comment_id, "")
                    if comment_text:
                        all_comments.append(comment_text)

                if custom_note.strip():
                    all_comments.append(custom_note.strip())

    # Append resubmittal comment at the very end if "Yes" was selected
    if st.session_state.wizard_resubmittal == "Yes":
        resubmittal_text = COMMENTS.get("BB-0045", "")
        if resubmittal_text:
            all_comments.append(resubmittal_text)

    return all_comments


def generate_word_document():
    """Generate a Word document with review comments"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('Engineering Plan Review', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # PROJECT INFORMATION
    # =========================================================================
    doc.add_heading('Project Information', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Review Type:', st.session_state.wizard_review_type or 'Not specified'),
        ('Permit Number:', st.session_state.wizard_permit_number or 'Not specified'),
        ('Address:', st.session_state.wizard_address or 'Not specified'),
        ('Reviewer:', st.session_state.wizard_reviewer or 'Not specified'),
        ('Review Date:', datetime.now(_CT).strftime('%Y-%m-%d %I:%M %p')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # =========================================================================
    # SUMMARY STATISTICS
    # =========================================================================
    doc.add_heading('Review Summary', level=1)
    
    yes_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "Yes")
    no_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "No")
    na_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "N/A")
    total = yes_count + no_count + na_count
    
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Total Items Reviewed:', str(total)),
        ('Compliant (Yes):', str(yes_count)),
        ('Issues Found (No):', str(no_count)),
        ('Not Applicable:', str(na_count)),
        ('Resubmittal Comment:', st.session_state.wizard_resubmittal),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # FULL CHECKLIST WITH STATUS AND COMMENTS
    # =========================================================================
    doc.add_heading('Plan Review Checklist', level=1)
    
    checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
    
    for section_id, section_data in checklist.items():
        section_heading = doc.add_heading(section_data["name"], level=2)
        
        for item in section_data["items"]:
            item_key = item["id"]
            status = st.session_state.wizard_checklist_state.get(item_key, "Not Reviewed")
            
            para = doc.add_paragraph()
            item_run = para.add_run(f"{item['id']} - {item['description']}")
            item_run.bold = False
            
            para.add_run("\n")
            status_run = para.add_run(f"Status: {status}")
            status_run.bold = True
            
            if status == "Yes":
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            elif status == "No":
                status_run.font.color.rgb = RGBColor(200, 0, 0)
            elif status == "N/A":
                status_run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                status_run.font.color.rgb = RGBColor(255, 165, 0)
            
            if status == "No":
                selected = st.session_state.wizard_selected_comments.get(item_key, [])
                custom_note = st.session_state.wizard_custom_notes.get(item_key, "")
                
                if selected or custom_note.strip():
                    para.add_run("\n")
                    comments_label = para.add_run("Comments:")
                    comments_label.bold = True
                    comments_label.font.color.rgb = RGBColor(0, 0, 128)
                    
                    for comment_id in selected:
                        comment_text = COMMENTS.get(comment_id, "")
                        if comment_text:
                            comment_para = doc.add_paragraph()
                            comment_para.paragraph_format.left_indent = Inches(0.5)
                            comment_run = comment_para.add_run(f"• [{comment_id}] {comment_text}")
                            comment_run.font.size = Pt(10)
                    
                    if custom_note.strip():
                        custom_para = doc.add_paragraph()
                        custom_para.paragraph_format.left_indent = Inches(0.5)
                        custom_run = custom_para.add_run(f"• [CUSTOM] {custom_note}")
                        custom_run.font.size = Pt(10)
                        custom_run.italic = True
        
        doc.add_paragraph()
    
    # =========================================================================
    # RESUBMITTAL STATUS IN DOCUMENT
    # =========================================================================
    if st.session_state.wizard_resubmittal == "Yes":
        doc.add_heading('Resubmittal', level=2)
        resub_para = doc.add_paragraph()
        resub_run = resub_para.add_run("Standard resubmittal comment included (BB-0045)")
        resub_run.bold = True
        resub_run.font.color.rgb = RGBColor(0, 0, 128)
    
    # =========================================================================
    # COMMENTS FOR COPY/PASTE (Only "No" items + resubmittal)
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Comments for Copy/Paste', level=1)
    
    intro_para = doc.add_paragraph()
    intro_para.add_run("Use the comments below for Bluebeam or permit system. ").italic = True
    intro_para.add_run("Only items marked 'No' with selected comments are included.").italic = True
    if st.session_state.wizard_resubmittal == "Yes":
        intro_para.add_run(" Resubmittal comment appended at end.").italic = True
    doc.add_paragraph()
    
    all_comments = collect_all_comments()
    
    if all_comments:
        for i, comment in enumerate(all_comments, 1):
            para = doc.add_paragraph()
            num_run = para.add_run(f"{i}. ")
            num_run.bold = True
            para.add_run(comment)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No comments to include - all items are compliant or N/A.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# LAMA CSV EXPORT
# =============================================================================

def generate_lama_csv():
    """
    Generate CSV for the LAMA Comment Uploader chrome extension.
    Format: Single column with header 'Comments', RFC 4180 quoting.
    Includes resubmittal comment at end if selected.
    """
    comments = collect_all_comments()
    if not comments:
        return None

    buffer = StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_ALL)
    writer.writerow(["Comments"])
    for comment in comments:
        writer.writerow([comment])

    return buffer.getvalue().encode('utf-8')


# =============================================================================
# BLUEBEAM BAX EXPORT (Bluebeam Markup Archive)
# =============================================================================

def _generate_annotation_id():
    """Generate a 16-character uppercase letter ID matching Bluebeam convention."""
    return ''.join(random.choices(string.ascii_uppercase, k=16))


def _pdf_escape(text):
    """Escape special characters for PDF string literals inside ()."""
    return (str(text)
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)"))


def _xml_escape(text):
    """Escape special characters for XML text content."""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;"))


def _build_annotation_raw(comment_text, reviewer, annot_id, rect, pdf_date):
    """
    Build the zlib-compressed, hex-encoded Raw field for a BAX annotation.
    Contains the full PDF annotation dictionary with styling:
      - Green (#008000) border and fill
      - 25% fill opacity
      - Helvetica 12pt black text
      - 0.75pt solid border
    """
    x1, y1, x2, y2 = rect
    pdf_text = _pdf_escape(comment_text)
    pdf_reviewer = _pdf_escape(reviewer)
    html_text = (comment_text
                 .replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;"))
    rc_text = _pdf_escape(html_text)

    raw_str = (
        '<</DA(0 0.5019608 0 rg /Helv 12 Tf)'
        '/DS(font: Helvetica 12pt; text-align:left; margin:0pt; '
        'line-height:13.8pt; color:#000000)'
        f'/TempBBox[{x1} {y1} {x2} {y2}]'
        '/FillOpacity 0.25'
        f'/T({pdf_reviewer})'
        f'/CreationDate({pdf_date})'
        '/RC(<?xml version="1.0"?>'
        '<body xmlns:xfa="http://www.xfa.org/schema/xfa-data/1.0/"'
        ' xfa:contentType="text/html"'
        ' xfa:APIVersion="BluebeamPDFRevu:2018"'
        ' xfa:spec="2.2.0"'
        ' style="font: Helvetica 12pt; text-align:left; margin:0pt; '
        'line-height:13.8pt; color:#000000"'
        ' xmlns="http://www.w3.org/1999/xhtml">'
        f'<p>{rc_text}</p></body>)'
        '/Subj(Engineering)'
        f'/NM({annot_id})'
        '/Subtype/FreeText'
        f'/Rect[{x1} {y1} {x2} {y2}]'
        f'/Contents({pdf_text})'
        '/F 4'
        '/C[0 0.5019608 0]'
        '/BS<</W 0.75/S/S/Type/Border>>'
        f'/M({pdf_date})>>'
    )

    compressed = zlib.compress(raw_str.encode('utf-8'))
    return compressed.hex()


def generate_bluebeam_bax():
    """
    Generate a Bluebeam BAX file with fully styled FreeText annotations.
    Includes resubmittal comment at end if selected.
    """
    comments = collect_all_comments()
    if not comments:
        return None

    reviewer = st.session_state.wizard_reviewer or "Engineering"
    now = datetime.now(_CT)
    iso_date = now.strftime("%Y-%m-%dT%H:%M:%S") + ".0000000Z"
    pdf_date = now.strftime("D:%Y%m%d%H%M%S") + "-06'00'"

    page_width, page_height = 612, 792
    box_width, box_height = 252, 108
    margin, gap = 36, 10
    current_y_top = page_height - margin
    x1 = margin

    annotation_blocks = []

    for i, comment in enumerate(comments):
        y2 = current_y_top
        y1 = current_y_top - box_height

        if y1 < margin and x1 == margin:
            x1 = margin + box_width + margin
            current_y_top = page_height - margin
            y2 = current_y_top
            y1 = current_y_top - box_height

        x2 = x1 + box_width
        rect = (x1, y1, x2, y2)
        annot_id = _generate_annotation_id()
        raw_hex = _build_annotation_raw(comment, reviewer, annot_id, rect, pdf_date)

        annotation_xml = (
            '    <Annotation>\n'
            '      <Page>1</Page>\n'
            f'      <Contents>{_xml_escape(comment)}</Contents>\n'
            f'      <ModDate>{iso_date}</ModDate>\n'
            '      <Color>#008000</Color>\n'
            '      <Type>FreeText</Type>\n'
            f'      <ID>{annot_id}</ID>\n'
            '      <TypeInternal>Bluebeam.PDF.Annotations.AnnotationFreeText'
            '</TypeInternal>\n'
            f'      <Raw>{raw_hex}</Raw>\n'
            f'      <Index>{i}</Index>\n'
            '      <Subject>Engineering</Subject>\n'
            f'      <CreationDate>{iso_date}</CreationDate>\n'
            f'      <Author>{_xml_escape(reviewer)}</Author>\n'
            '    </Annotation>'
        )

        annotation_blocks.append(annotation_xml)
        current_y_top = y1 - gap

    annotations_str = '\n'.join(annotation_blocks)

    bax_content = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<Document Version="1">\n'
        '  <Page Index="0">\n'
        '    <Label>1</Label>\n'
        f'    <Width>{page_width}</Width>\n'
        f'    <Height>{page_height}</Height>\n'
        f'{annotations_str}\n'
        '  </Page>\n'
        '</Document>'
    )

    bax_crlf = bax_content.replace('\n', '\r\n')
    return b'\xef\xbb\xbf' + bax_crlf.encode('utf-8')


@st.fragment
def _render_checklist():
    """
    Checklist + resubmittal section as a st.fragment.

    st.fragment causes ONLY this section to rerender when a Yes/No/NA button
    is clicked. The page header, sidebar, CSS injection, and Step 3 summary
    do NOT rerender — button clicks feel instantaneous.
    """
    review_type = st.session_state.wizard_review_type
    if not review_type:
        return

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown(f"<div class='bw-step-heading'>Step 2 — {review_type} Checklist</div>", unsafe_allow_html=True)

    checklist     = get_checklist_for_review_type(review_type)
    total_items   = sum(len(s["items"]) for s in checklist.values())
    completed_items = len(st.session_state.wizard_checklist_state)

    st.progress(completed_items / total_items if total_items > 0 else 0)
    st.caption(f"Progress: {completed_items} of {total_items} items reviewed")

    # ── Sections as collapsed expanders ───────────────────────────────────
    # wizard_open_section tracks which section the user last clicked in,
    # so the expander stays open after a button press.
    if "wizard_open_section" not in st.session_state:
        st.session_state.wizard_open_section = list(checklist.keys())[0]

    # Sequential display counters — these control what numbers the user SEES.
    # The internal item["id"] values (e.g. "6.3", "16.2") never change because
    # they are used as session state keys, comment lookups, and export labels.
    # We simply replace the hard-coded numbers in the display with a clean
    # 1, 2, 3… for sections and 1.1, 1.2… for items so there are never gaps
    # caused by deleted items or permit-type filtering.
    display_section_num = 0

    for section_id, section_data in checklist.items():
        section_items = section_data["items"]
        section_done  = sum(
            1 for it in section_items
            if it["id"] in st.session_state.wizard_checklist_state
        )
        section_total = len(section_items)
        all_done      = section_done == section_total

        # Increment the visible section number for each section that appears
        display_section_num += 1

        # Strip the original hard-coded number from the section name
        # e.g. "6. Retaining Walls" → "Retaining Walls", then prepend display num
        raw_name = section_data["name"]
        # Remove leading digits and period/dot if present (e.g. "6. " or "16. ")
        clean_name = re.sub(r'^\d+\.\s*', '', raw_name)
        display_section_name = f"{display_section_num}. {clean_name}"

        status_icon     = "✅" if all_done else "📋"
        expander_label  = (
            f"{status_icon} {display_section_name}  "
            f"({section_done}/{section_total} reviewed)"
        )

        default_open = (section_id == st.session_state.wizard_open_section)

        with st.expander(expander_label, expanded=default_open):
            # Sequential item counter resets for each section: 1.1, 1.2, 1.3…
            display_item_num = 0

            for item in section_items:
                item_key       = item["id"]
                current_status = st.session_state.wizard_checklist_state.get(item_key, "—")

                display_item_num += 1
                display_item_label = f"{display_section_num}.{display_item_num}"

                # ── Item row: description + Yes / No / N/A buttons ────────────
                desc_col, yes_col, no_col, na_col = st.columns([5, 1, 1, 1])

                with desc_col:
                    st.markdown(
                        f"<div style='padding:0.4rem 0;font-size:0.93rem;'>"
                        f"<strong style='color:#22427C'>{display_item_label}</strong>"
                        f" — {item['description']}</div>",
                        unsafe_allow_html=True
                    )

                with yes_col:
                    def _set_yes(k=item_key, s=section_id):
                        st.session_state.wizard_checklist_state[k] = "Yes"
                        st.session_state.wizard_selected_comments.pop(k, None)
                        st.session_state.wizard_open_section = s
                    yes_type = "primary" if current_status == "Yes" else "secondary"
                    st.button("✓ Yes", key=f"yes_{item_key}", type=yes_type,
                              use_container_width=True, on_click=_set_yes)

                with no_col:
                    def _set_no(k=item_key, s=section_id):
                        st.session_state.wizard_checklist_state[k] = "No"
                        st.session_state.wizard_open_section = s
                    no_type = "primary" if current_status == "No" else "secondary"
                    st.button("✗ No", key=f"no_{item_key}", type=no_type,
                              use_container_width=True, on_click=_set_no)

                with na_col:
                    def _set_na(k=item_key, s=section_id):
                        st.session_state.wizard_checklist_state[k] = "N/A"
                        st.session_state.wizard_selected_comments.pop(k, None)
                        st.session_state.wizard_open_section = s
                    na_type = "primary" if current_status == "N/A" else "secondary"
                    st.button("N/A", key=f"na_{item_key}", type=na_type,
                              use_container_width=True, on_click=_set_na)

                # ── Comment panel (only when No selected) ─────────────────────
                if st.session_state.wizard_checklist_state.get(item_key) == "No":
                    with st.container():
                        st.markdown('<div class="bw-comment-box">', unsafe_allow_html=True)
                        st.markdown("**📝 Select applicable comments:**")

                        comment_ids = item.get("comment_ids", [])
                        if comment_ids:
                            if item_key not in st.session_state.wizard_selected_comments:
                                st.session_state.wizard_selected_comments[item_key] = []

                            for comment_id in comment_ids:
                                comment_text = COMMENTS.get(comment_id, "Comment not found")
                                display_text = comment_text[:150] + "..." if len(comment_text) > 150 else comment_text
                                is_selected  = comment_id in st.session_state.wizard_selected_comments[item_key]

                                if st.checkbox(
                                    f"**{comment_id}**: {display_text}",
                                    value=is_selected,
                                    key=f"comment_{item_key}_{comment_id}"
                                ):
                                    if comment_id not in st.session_state.wizard_selected_comments[item_key]:
                                        st.session_state.wizard_selected_comments[item_key].append(comment_id)
                                else:
                                    if comment_id in st.session_state.wizard_selected_comments[item_key]:
                                        st.session_state.wizard_selected_comments[item_key].remove(comment_id)

                                # Nested expanders are illegal in Streamlit —
                                # show the full text as a small caption instead.
                                if len(comment_text) > 150:
                                    st.caption(comment_text)

                        st.markdown("**✏️ Custom notes (optional):**")
                        custom_note = st.text_area(
                            "Custom note",
                            value=st.session_state.wizard_custom_notes.get(item_key, ""),
                            key=f"custom_{item_key}",
                            height=80,
                            label_visibility="collapsed",
                            placeholder="Add any additional comments specific to this review..."
                        )
                        st.session_state.wizard_custom_notes[item_key] = custom_note
                        st.markdown('</div>', unsafe_allow_html=True)

                st.markdown("---")


def _tc_init():
    """
    Initialize all tc_ session state keys before widgets render.
    Uses setdefault() throughout — this is the only safe initialization
    pattern for widget-bound keys in Streamlit >= 1.30.
    Direct assignment (st.session_state[key] = value) raises
    StreamlitAPIException when the key is already claimed by a widget.
    setdefault() is a no-op when the key already exists, so it never
    conflicts with a widget that has already written its value.
    """
    str_keys = [
        "tc_case_num", "tc_street_class", "tc_street_name",
        "tc_street_name_input", "tc_street_segment", "tc_seg_length",
        "tc_app_type", "tc_petitioner_name", "tc_petitioner_contact",
        "tc_hoa_status", "tc_eligible_res", "tc_sigs_received", "tc_init_pct",
        "tc_problem_desc", "tc_data_85th", "tc_data_limit", "tc_data_adt",
        "tc_data_crashes", "tc_data_cutthru", "tc_speed_excess",
        "tc_speed_excess_raw", "tc_data_school_route", "tc_data_sidewalk_status",
        "tc_data_notes", "tc_local_adt", "tc_local_adt_proj", "tc_local_width",
        "tc_local_grade", "tc_local_spd_limit",
        "tc_t1_notes", "tc_pet2_total", "tc_pet2_yes",
        "tc_pet2_pct",
        "tc_cost_direct",
        "tc_cost_contingency", "tc_cost_total", "tc_cost_resident",
        "tc_cost_city", "tc_cost_notes",
        "tc_public_meeting_notes", "tc_board_res_num",
        "tc_staff_rec_notes", "tc_final_notes",
    ]
    # Date keys — stored as datetime.date or None.
    # st.date_input(value=None) renders as a blank optional field (Streamlit 1.38+).
    # The report builder reads formatted strings; the date widgets auto-format on render.
    date_keys = [
        "tc_app_date",
        "tc_t1_date",
        "tc_t1_review_date",
        "tc_pet2_mail",
        "tc_moratorium_start",
        "tc_public_meeting_date",
        "tc_board_date",
    ]
    bool_keys = [
        "tc_c_hoa_gov", "tc_c_hoa_req", "tc_c_init_pet", "tc_c_pet_format",
        "tc_c_not_private", "tc_c_not_emergency", "tc_c_collector_list",
        "tc_c_data_speed", "tc_c_data_vol", "tc_c_data_crash",
        "tc_c_data_school", "tc_c_data_sidewalk",
        "tc_c_coll_speed_data", "tc_c_coll_2lanes", "tc_c_coll_termini",
        "tc_c_coll_study_vol", "tc_c_coll_study_speed", "tc_c_coll_study_crash",
        "tc_c_coll_study_sidewalk", "tc_c_coll_study_school", "tc_c_coll_study_t2",
        "tc_c_local_lanewidth", "tc_c_local_grade", "tc_c_local_speedlimit",
        "tc_c_local_notarterial", "tc_c_local_cutthru", "tc_c_local_connection",
        "tc_c_hump_spacing", "tc_c_hump_clearance", "tc_c_hump_dims",
        "tc_c_hump_signage", "tc_c_hump_warn", "tc_c_hump_markings",
        "tc_c_hump_drainage", "tc_c_t1_study", "tc_c_t1_petitioner",
        "tc_c_t1_implemented", "tc_c_t1_effective", "tc_c_t1_ineffective",
        "tc_c_t2_validate", "tc_c_t2_trafficeng", "tc_c_t2_sep_petition",
        "tc_c_pet2_mailed", "tc_c_pet2_hoa", "tc_c_pet2_nonresp", "tc_pet2_ext",
        "tc_c_costshare_agree", "tc_c_hoa_letter", "tc_c_cost_payment",
        "tc_c_public_meeting", "tc_c_public_conducted", "tc_c_staff_rec",
        "tc_c_board_res", "tc_c_board_action", "tc_c_design_final",
        "tc_c_payment_rcvd", "tc_c_installed", "tc_c_archived",
        "tc_c_leftover_funds",
    ]
    for k in str_keys:
        st.session_state.setdefault(k, "")
    for k in date_keys:
        st.session_state.setdefault(k, None)
    for k in bool_keys:
        st.session_state.setdefault(k, False)
    st.session_state.setdefault("tc_t2_strategies", [])
    st.session_state.setdefault("tc_total_score", 0)
    for crit in SCORING_CRITERIA:
        st.session_state.setdefault(f"tc_score_{crit['id']}", 0)



def render_traffic_calming_wizard():
    """
    Traffic Calming Application Review form.
    Uses key="tc_X" on every widget so Streamlit manages session_state automatically.
    No left-hand st.session_state assignments on widget calls (causes StreamlitAPIException).
    _tc_init() pre-populates all keys with correct types before any widget renders.
    """
    _tc_init()

    st.markdown(
        "<div class='bw-step-heading'>Traffic Calming Application Review</div>",
        unsafe_allow_html=True,
    )
    st.caption("Policy: Resolution 2026-12 · City of Brentwood, TN · Adopted 02/09/2026")

    def safe_idx(options, value):
        try:
            return options.index(value)
        except ValueError:
            return 0

    def tc_date(label, key, help_text=None, override_value=None):
        """
        Render a date_input that stores a datetime.date or None.
        value=None renders as a blank optional field (Streamlit 1.38+).
        Also writes a formatted string to key + '_str' so the report builder
        can read it without needing to handle date objects.

        override_value: pass a datetime.date to pre-fill the picker without
        writing to session state directly (avoids the Streamlit warning about
        setting a widget value via both value= and session state).
        Only used when the field has not yet been set by the user.
        """
        import datetime as dt
        # Determine the value to show: prefer what's already in session state
        # (user's own selection), then fall back to override_value, then None.
        current = st.session_state.get(key, None)
        display_value = current if current is not None else override_value
        val = st.date_input(
            label,
            value=display_value,
            key=key,
            help=help_text,
            format="MM/DD/YYYY",
        )
        # Write formatted string for the report builder
        date_str = val.strftime("%m/%d/%Y") if isinstance(val, dt.date) else ""
        st.session_state[key + "_str"] = date_str
        return val

    # =========================================================================
    # SAVE / LOAD PROGRESS
    # =========================================================================

    def _tc_filename():
        """Build a safe filename from case number + street name."""
        import re
        case   = st.session_state.get("tc_case_num", "") or ""
        street = st.session_state.get("tc_street_name", "") or ""
        parts  = [p for p in [case, street] if p]
        raw    = "_".join(parts) if parts else "TC_progress"
        # Replace spaces and common punctuation with underscores, strip the rest
        safe = re.sub(r"[\s/\\:*?\"<>|]+", "_", raw)
        safe = re.sub(r"[^\w\-]", "", safe)
        safe = re.sub(r"_+", "_", safe).strip("_")
        return f"{safe}_progress.json"

    def _save_tc_state() -> bytes:
        """
        Serialize all tc_ session state keys to JSON bytes.
        datetime.date objects are stored as ISO strings with a __date__ marker
        so they can be round-tripped correctly on load.
        """
        import json
        import datetime as dt
        payload = {}
        for k, v in st.session_state.items():
            if not k.startswith("tc_"):
                continue
            if isinstance(v, dt.date):
                payload[k] = {"__date__": v.isoformat()}
            elif isinstance(v, (str, bool, int, float, list, type(None))):
                payload[k] = v
            # skip anything else (e.g. widget-internal objects)
        return json.dumps(payload, indent=2).encode("utf-8")

    def _load_tc_state(uploaded_file):
        """
        Restore tc_ session state from an uploaded JSON file.
        Converts __date__ markers back to datetime.date objects.
        Returns (success: bool, message: str).
        """
        import json
        import datetime as dt
        try:
            raw = uploaded_file.read()
            payload = json.loads(raw)
        except Exception as e:
            return False, f"Could not read file: {e}"

        if not isinstance(payload, dict):
            return False, "Invalid file format — expected a JSON object."

        # Clear existing tc_ keys first so stale values don't linger
        for k in list(st.session_state.keys()):
            if k.startswith("tc_"):
                del st.session_state[k]

        # Restore values with correct types
        loaded = 0
        for k, v in payload.items():
            if not k.startswith("tc_"):
                continue
            if isinstance(v, dict) and "__date__" in v:
                try:
                    st.session_state[k] = dt.date.fromisoformat(v["__date__"])
                    loaded += 1
                except Exception:
                    pass  # skip malformed dates
            else:
                st.session_state[k] = v
                loaded += 1

        return True, f"Loaded {loaded} fields successfully."

    # ── Save / Load bar ───────────────────────────────────────────────────────
    with st.container():
        save_col, load_col = st.columns([1, 2])

        with save_col:
            if st.button("💾 Save Progress", use_container_width=True,
                         key="btn_tc_save",
                         help="Download current form state as a JSON file to resume later"):
                _data = _save_tc_state()
                _fname = _tc_filename()
                st.download_button(
                    label=f"⬇ Download  {_fname}",
                    data=_data,
                    file_name=_fname,
                    mime="application/json",
                    use_container_width=True,
                    key="btn_tc_save_dl",
                )

        with load_col:
            uploaded_json = st.file_uploader(
                "Load saved progress (.json)",
                type=["json"],
                key="tc_load_uploader",
                help="Upload a previously saved progress file to restore all form fields",
                label_visibility="collapsed",
            )
            if uploaded_json is not None:
                ok, msg = _load_tc_state(uploaded_json)
                if ok:
                    st.success(f"✓ {msg}")
                    st.rerun()
                else:
                    st.error(f"Load failed: {msg}")

    st.divider()

    # =========================================================================
    # SECTION I: ADMINISTRATIVE REVIEW
    # =========================================================================
    with st.expander("I. Administrative Review", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.text_input("Case / File Number", key="tc_case_num",
                          placeholder="e.g. TC-2026-001")
        with col2:
            tc_date("Application Date", key="tc_app_date")

        st.markdown("---")
        st.markdown("**Step 1 — Street Classification**")

        class_opts = ["-- Select --"] + list(STREET_CLASSIFICATIONS.keys())
        prev_class = st.session_state.get("tc_street_class", "")
        chosen_class = st.selectbox(
            "Street Classification",
            options=class_opts,
            index=safe_idx(class_opts, prev_class),
            key="tc_street_class_sel",
            help="Arterial/Collector show City-designated street lists. Local = free text.",
        )
        street_class = "" if chosen_class == "-- Select --" else chosen_class

        # When classification changes, store the new class and trigger a rerun
        # so old widget keys are no longer in the widget tree before we clear them.
        # Never write directly to tc_street_name_input — that's a widget-bound key.
        if street_class != prev_class:
            st.session_state["tc_street_class"] = street_class
            st.session_state["tc_street_name"]  = ""
            # Remove the old Local text_input widget key from session state
            # ONLY if it isn't currently the active widget (i.e. we're switching away).
            if "tc_street_name_input" in st.session_state and street_class != "Local Residential Street":
                del st.session_state["tc_street_name_input"]
            st.rerun()
        else:
            st.session_state["tc_street_class"] = street_class

        if street_class == "Collector Street":
            st.success("✓ **Collector Street** — Traffic Calming Policy (Part V). 100% City-funded.")
        elif street_class == "Local Residential Street":
            st.info("ℹ **Local Street** — Speed Hump Policy (Part VII). Residents pay 60% of costs.")
        elif street_class == "Arterial Street":
            st.warning("⚠ **Arterial Street** — Standard calming policy does NOT apply. Engineering study required.")

        st.markdown("**Step 2 — Street Name**")

        if street_class == "Arterial Street":
            art_opts = ["-- Select Arterial --"] + ARTERIAL_STREETS
            chosen_art = st.selectbox(
                "Arterial Street (Municipal Code — Arterial Designation)",
                options=art_opts,
                index=safe_idx(art_opts, st.session_state.get("tc_street_name", "")),
                key="tc_street_name_art",
            )
            # Store into a non-widget-bound data key
            new_name = "" if chosen_art == "-- Select Arterial --" else chosen_art
            if new_name != st.session_state.get("tc_street_name", ""):
                st.session_state["tc_street_name"] = new_name

        elif street_class == "Collector Street":
            col_opts = ["-- Select Collector --"] + COLLECTOR_STREETS
            chosen_col = st.selectbox(
                "Collector Street (Municipal Code — Collector Designation)",
                options=col_opts,
                index=safe_idx(col_opts, st.session_state.get("tc_street_name", "")),
                key="tc_street_name_col",
            )
            new_name = "" if chosen_col == "-- Select Collector --" else chosen_col
            if new_name != st.session_state.get("tc_street_name", ""):
                st.session_state["tc_street_name"] = new_name

        elif street_class == "Local Residential Street":
            # tc_street_name_input is the widget key — never write to it directly.
            # Streamlit manages its value. We read it back into tc_street_name
            # only when the value has changed, using a safe get (not assignment).
            st.text_input(
                "Local Street Name (enter manually)",
                key="tc_street_name_input",
                placeholder="e.g. Oakwood Court",
            )
            # Sync widget value → data store using a non-widget key
            current_input = st.session_state.get("tc_street_name_input", "")
            if current_input != st.session_state.get("tc_street_name", ""):
                st.session_state["tc_street_name"] = current_input

        else:
            st.text_input(
                "Street Name (select classification above first)",
                value="",
                disabled=True,
                key="tc_street_name_placeholder",
            )

        st.markdown("**Step 3 - Street Segment / Limits**")
        st.text_input("Segment Description / Limits  (e.g. from Oak Drive to Elm Street)",
                      key="tc_street_segment", placeholder="From ________ to ________")
        st.text_input("Approximate Segment Length (ft)  [Part V-a; Part VII]",
                      key="tc_seg_length", placeholder="ft")

        st.markdown("---")
        st.markdown("**Step 4 - Application Type**")
        type_opts = ["-- Select --"] + list(APPLICATION_TYPES.keys())
        chosen_type = st.selectbox(
            "Application Type",
            options=type_opts,
            index=safe_idx(type_opts, st.session_state.get("tc_app_type", "")),
            key="tc_app_type_sel",
        )
        st.session_state["tc_app_type"] = "" if chosen_type == "-- Select --" else chosen_type

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            st.text_input("Petitioner Name / Organization", key="tc_petitioner_name")
        with col2:
            st.text_input("Petitioner Contact (Phone / Email)", key="tc_petitioner_contact")

        st.markdown("**HOA Prerequisite** *(Part V - Procedures; Part VII)*")
        hoa_opts = [
            "-- Select --",
            "HOA submitting request to City",
            "HOA denied - resident proceeding directly",
            "HOA inaction after 90 days - resident proceeding",
            "No HOA - resident petition permitted",
        ]
        chosen_hoa = st.selectbox(
            "HOA Determination",
            options=hoa_opts,
            index=safe_idx(hoa_opts, st.session_state.get("tc_hoa_status", "")),
            key="tc_hoa_status_sel",
        )
        st.session_state["tc_hoa_status"] = "" if chosen_hoa == "-- Select --" else chosen_hoa

        st.checkbox("HOA governance confirmed  [Part V / VII]", key="tc_c_hoa_gov")
        st.checkbox("HOA written request submitted - or bypass justified (denied / 90-day inaction)  [Part V / VII]",
                    key="tc_c_hoa_req")

        st.markdown("**Initial Support Petition** - more than 50% of homes within 600 ft; one signature/name per residence; eligibility ends 100 ft before a stop sign  [Part V / VII]")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_input("Eligible Residences (#)", key="tc_eligible_res")
        with col2:
            st.text_input("Signatures Received (#)", key="tc_sigs_received")
        with col3:
            try:
                pct_val = (float(st.session_state.get("tc_sigs_received") or 0)
                           / float(st.session_state.get("tc_eligible_res") or 1) * 100)
                pct_str = f"{pct_val:.1f}%"
                st.session_state["tc_init_pct"] = pct_str
                color = "green" if pct_val > 50 else "red"
                st.markdown(f"**Initial Support:** :{color}[{pct_str}]")
                if pct_val <= 50:
                    st.caption("Must exceed 50% to initiate study")
            except Exception:
                st.markdown("**Initial Support:** -")

        st.checkbox("Initial more-than-50% support petition obtained  [Part V / VII]", key="tc_c_init_pet")
        st.checkbox("One signature and printed name per residence confirmed  [Part V]", key="tc_c_pet_format")
        st.text_area("Description of Perceived Problem / Safety Concern",
                     key="tc_problem_desc", height=80)

    # =========================================================================
    # SECTION II: ELIGIBILITY
    # =========================================================================
    with st.expander("II. Classification Eligibility Confirmation", expanded=False):
        st.checkbox("Public street confirmed (not a private street in a gated community)  [Part V / VII]",
                    key="tc_c_not_private")
        st.checkbox("Confirmed NOT a designated primary emergency route  [Part V]",
                    key="tc_c_not_emergency")
        if street_class == "Collector Street":
            st.checkbox("Collector: verified on City identified residential collector street list  [Part V]",
                        key="tc_c_collector_list")
            st.caption("Residential collector list per Resolution 2026-12 Part V: Arrowhead Dr, Belle Rive Dr, "
                       "Bluff Rd, Carriage Hills Dr, Charity Dr, Concord Pass, Gen. George Patton Dr, "
                       "Gordon Petty Rd, Johnson Chapel Rd W, Jones Pkwy, Knox Valley Dr, Lipscomb Dr, "
                       "Manley Ln, McGavock Rd, Pinkerton Rd, Stanfield Rd, Steeplechase Dr, "
                       "Sunset Rd (N of Concord Rd), Walnut Hills Dr.")

    # =========================================================================
    # SECTION III: DATA COLLECTION
    # =========================================================================
    with st.expander("III. Data Collection & Field Review", expanded=False):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_input("85th Percentile Speed (mph)  [Part V-a / VII]", key="tc_data_85th")
            st.text_input("ADT (vehicles/day)  [Part V-a / VII]", key="tc_data_adt")
        with col2:
            st.text_input("Posted Speed Limit (mph)", key="tc_data_limit")
            st.text_input("Crash Count (12 months)  [Parts I-II]", key="tc_data_crashes")
        with col3:
            try:
                excess_val = (float(st.session_state.get("tc_data_85th") or 0)
                              - float(st.session_state.get("tc_data_limit") or 0))
                excess_str = f"{excess_val:.1f} mph over limit"
                st.session_state["tc_speed_excess"] = excess_str
                st.session_state["tc_speed_excess_raw"] = str(excess_val)
                color = "green" if excess_val >= 8 else "red"
                st.markdown(f"**Speed Excess:** :{color}[{excess_str}]")
                if street_class == "Collector Street" and excess_val < 8:
                    st.caption("Collector criterion: 8 mph over limit required [Part V-a]")
            except Exception:
                st.markdown("**Speed Excess:** -")
            st.text_input("Cut-Through Traffic % [Part VII - 35% or more = cut-through]",
                          key="tc_data_cutthru")

        col1, col2 = st.columns(2)
        with col1:
            school_opts = ["-- Select --", "Yes - street is on a school walking route",
                           "No - not a school walking route"]
            chosen_sch = st.selectbox(
                "School Walking Route?  [Part V-b-3: 10 pts if yes]",
                options=school_opts,
                index=safe_idx(school_opts, st.session_state.get("tc_data_school_route", "")),
                key="tc_data_school_route_sel",
            )
            st.session_state["tc_data_school_route"] = "" if chosen_sch == "-- Select --" else chosen_sch
        with col2:
            swlk_opts = ["-- Select --", "No continuous sidewalk", "Continuous sidewalk exists"]
            chosen_swlk = st.selectbox(
                "Continuous Sidewalk?  [Part V-b-3: 10 pts if no sidewalk]",
                options=swlk_opts,
                index=safe_idx(swlk_opts, st.session_state.get("tc_data_sidewalk_status", "")),
                key="tc_data_sidewalk_status_sel",
            )
            st.session_state["tc_data_sidewalk_status"] = "" if chosen_swlk == "-- Select --" else chosen_swlk

        st.checkbox("Speed study completed (24-hr weekday minimum for collectors)  [Part V-a]",
                    key="tc_c_data_speed")
        st.checkbox("Traffic count / ADT collected  [Part V-a / VII]", key="tc_c_data_vol")
        st.checkbox("Crash history reviewed (12 months minimum)  [Parts I-II]", key="tc_c_data_crash")
        st.checkbox("School route status confirmed  [Part V-b-3]", key="tc_c_data_school")
        st.checkbox("Continuous sidewalk presence/absence confirmed  [Part V-b-3]", key="tc_c_data_sidewalk")
        st.text_area("Data Collection Notes / Summary", key="tc_data_notes", height=80)

    # =========================================================================
    # SECTION IV: STREET-TYPE SPECIFIC
    # =========================================================================
    if street_class == "Collector Street":
        with st.expander("IV. Traffic Calming - Collector Street Criteria (Part V)", expanded=False):
            st.markdown("**Roadway Eligibility Criteria - All must be met [Part V-a]**")
            try:
                adt_ok = float(st.session_state.get("tc_data_adt") or 0) >= 500
                len_ok = float(st.session_state.get("tc_seg_length") or 0) >= 800
                spd_ok = float(str(st.session_state.get("tc_speed_excess_raw") or 0)) >= 8
                if not adt_ok:
                    st.error(f"ADT must be 500 or more vpd - entered: {st.session_state.get('tc_data_adt') or '?'}  [Part V-a Volume]")
                if not len_ok:
                    st.error(f"Segment must be 800 ft or more - entered: {st.session_state.get('tc_seg_length') or '?'} ft  [Part V-a Other Criteria]")
                if not spd_ok:
                    st.error(f"85th percentile must exceed limit by 8 mph or more  [Part V-a Speed]")
                if adt_ok and len_ok and spd_ok:
                    st.success("All three roadway criteria met - eligible to proceed with study.")
            except Exception:
                pass

            st.checkbox("24-hr weekday speed data collected  [Part V-a Speed]", key="tc_c_coll_speed_data")
            st.checkbox("Street has no more than two traffic lanes (one each direction)  [Part V-a]",
                        key="tc_c_coll_2lanes")
            st.checkbox("Logical termini for calming treatment identifiable  [Part V-a]", key="tc_c_coll_termini")
            st.markdown("**Engineering Study Contents Required [Part V-b-1]**")
            st.checkbox("Study includes traffic volume analysis", key="tc_c_coll_study_vol")
            st.checkbox("Study includes traffic speed analysis", key="tc_c_coll_study_speed")
            st.checkbox("Study includes accident history for subject segment", key="tc_c_coll_study_crash")
            st.checkbox("Study notes presence or absence of sidewalks", key="tc_c_coll_study_sidewalk")
            st.checkbox("Study addresses whether street is on a school walking route  [Part V-b-1 / V-b-3]",
                        key="tc_c_coll_study_school")
            st.checkbox("Study outlines applicable Tier 2 strategies for this location  [Part V-b-1]",
                        key="tc_c_coll_study_t2")

    elif street_class == "Local Residential Street":
        with st.expander("IV. Speed Hump Eligibility - Local Street (Part VII)", expanded=False):
            st.markdown("**Street Eligibility Criteria - All must be met [Part VII]**")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.text_input("Current ADT (vpd) - 500 to 2500 required", key="tc_local_adt")
                st.text_input("Street Width (ft) - less than 30 ft required", key="tc_local_width")
            with col2:
                st.text_input("Projected Full-Build ADT - 2500 or less required", key="tc_local_adt_proj")
                st.text_input("Street Grade (%) - 6% or less required", key="tc_local_grade")
            with col3:
                st.text_input("Posted Speed Limit (mph) - 30 mph or less required", key="tc_local_spd_limit")
                try:
                    adt_v = float(st.session_state.get("tc_local_adt") or 0)
                    if adt_v > 0:
                        if adt_v < 500 or adt_v > 2500:
                            st.error(f"ADT {adt_v:.0f} outside 500-2500 range  [Part VII]")
                        else:
                            st.success(f"ADT {adt_v:.0f} within range")
                except Exception:
                    pass

            st.checkbox("Two-lane street less than 30 ft wide confirmed  [Part VII]", key="tc_c_local_lanewidth")
            st.checkbox("Grade does not exceed 6%  [Part VII]", key="tc_c_local_grade")
            st.checkbox("Posted speed limit is 30 mph or less  [Part VII]", key="tc_c_local_speedlimit")
            st.checkbox("Confirmed NOT an arterial or collector street (speed humps prohibited on those)  [Part VII]",
                        key="tc_c_local_notarterial")
            st.checkbox("Street has identified cut-through or speeding problem confirmed by data  [Part VII]",
                        key="tc_c_local_cutthru")
            st.caption("Cut-through = 35% or more of traffic does not originate or terminate in subdivision. "
                       "Speeding = 85th percentile speed exceeds posted limit.")
            st.checkbox("Street provides connection between arterial/collector streets or subdivision pass-through  [Part VII]",
                        key="tc_c_local_connection")
            st.markdown("**Design Requirements - verified prior to installation [Part VII]**")
            st.checkbox("Minimum two humps proposed; spacing 300 to 600 ft apart  [Part VII]",
                        key="tc_c_hump_spacing")
            st.checkbox("Each hump 200 ft or more from any intersection and from any curve with radius 150 ft or less  [Part VII]",
                        key="tc_c_hump_clearance")
            st.checkbox("Hump max height 3 to 4 inches; travel length 12 ft per standard dimensions  [Part VII]",
                        key="tc_c_hump_dims")
            st.checkbox("Regulatory Residential Speed Control District signs installed (24x24, black on white)  [Part VII]",
                        key="tc_c_hump_signage")
            st.checkbox("MUTCD advance warning signs: 30x30 SPEED HUMPS sign plus 18x18 15 MPH advisory plate, about 125 ft before first hump  [Part VII]",
                        key="tc_c_hump_warn")
            st.checkbox("Double yellow centerline continued across all humps; 12-in white stripes at 6-inch O.C. per standard details  [Part VII]",
                        key="tc_c_hump_markings")
            st.checkbox("All hump locations reviewed by City Engineer - drainage adequately accommodated  [Part VII]",
                        key="tc_c_hump_drainage")

    # =========================================================================
    # SECTION V: TIER 1
    # =========================================================================
    with st.expander("V. Tier 1 - Non-Construction Strategies (Part V-b-1)", expanded=False):
        if street_class == "Local Residential Street":
            st.info("Less dramatic measures such as signs and striping must be evaluated first. "
                    "Reevaluate 6 months after installation before final speed hump decision.  [Part VII]")
        col1, col2 = st.columns(2)
        with col1:
            tc_date("Tier 1 Implementation Date", key="tc_t1_date")
        with col2:
            # Six-month review date auto-calculated from implementation date.
            # Passed as override_value so Streamlit never sees a session state
            # write conflict. User can select a different date to override.
            _t1_val = st.session_state.get("tc_t1_date")
            _auto_review = None
            if _t1_val:
                import datetime as _dt
                import calendar as _cal
                _m = _t1_val.month + 6
                _y = _t1_val.year + (_m - 1) // 12
                _m = (_m - 1) % 12 + 1
                _d = min(_t1_val.day, _cal.monthrange(_y, _m)[1])
                _auto_review = _dt.date(_y, _m, _d)
            tc_date("Six-Month Effectiveness Review Date", key="tc_t1_review_date",
                    help_text="Auto-calculated as 6 months after implementation date. Select a different date to override.",
                    override_value=_auto_review)
        st.checkbox("Study recommendation includes one or more Tier 1 strategies  [Part V-b-1]",
                    key="tc_c_t1_study")
        st.checkbox("Staff met with petitioner to outline study recommendations  [Part V-b-1]",
                    key="tc_c_t1_petitioner")
        st.checkbox("Tier 1 strategy implemented  [Part V-b-1]", key="tc_c_t1_implemented")
        col1, col2 = st.columns(2)
        with col1:
            st.checkbox("Tier 1 determined EFFECTIVE after 6 months - no further action  [Part V-b-1]",
                        key="tc_c_t1_effective")
        with col2:
            st.checkbox("Tier 1 determined INEFFECTIVE after 6 months - Tier 2 requested  [Part V-b-1]",
                        key="tc_c_t1_ineffective")
        if st.session_state.get("tc_c_t1_effective") and st.session_state.get("tc_c_t1_ineffective"):
            st.warning("Cannot mark both effective and ineffective - please uncheck one.")
        st.text_area("Tier 1 Notes (strategies applied, effectiveness findings)",
                     key="tc_t1_notes", height=80)

    # =========================================================================
    # SECTION VI: TIER 2 / SECOND PETITION
    # =========================================================================
    with st.expander("VI. Tier 2 / Speed Humps - Second Petition (Parts V-b-2; VII)", expanded=False):
        if street_class == "Collector Street":
            st.caption("Speed humps are NOT eligible on designated collector roads.  [Part V-b Tier 2 Note]")

        st.markdown("**Tier 2 Strategies Proposed [Part V-b Tier 2]**")
        t2_selected = []
        for strat, cite in TIER2_STRATEGIES:
            safe_key = ("tc_t2_"
                        + strat[:25].replace(" ", "_").replace("/", "_")
                                    .replace(",", "").replace("-", "").lower())
            if safe_key not in st.session_state:
                st.session_state[safe_key] = False
            st.checkbox(f"{strat}  [{cite}]", key=safe_key)
            if st.session_state.get(safe_key):
                t2_selected.append(strat)
        st.session_state["tc_t2_strategies"] = t2_selected

        st.markdown("**Required Review Steps [Part V-b-2]**")
        st.checkbox("City conducted second study validating Tier 1 was ineffective  [Part V-b-2]",
                    key="tc_c_t2_validate")
        st.checkbox("Tier 2 strategy reviewed by City Traffic Engineer before recommendation  [Part V-b-2]",
                    key="tc_c_t2_trafficeng")
        st.checkbox("Separate petition prepared for each proposed improvement  [Part V-b-2]",
                    key="tc_c_t2_sep_petition")
        st.caption("For a series of improvements such as speed tables or circles, one petition per improvement required.")

        st.markdown("**Second-Round Petition Process [Part V-b-2; Part VII]**")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_input("Total Households - Affected Area (within 600 ft)  [Part V-b-2 / VII]",
                          key="tc_pet2_total")
            tc_date("Petition Mail Date  [Part V-b-2 / VII]", key="tc_pet2_mail")
        with col2:
            st.text_input("Yes Votes Received  [Part V-b-2 / VII]", key="tc_pet2_yes")
            # 45-day deadline auto-calculated from petition mail date + extension
            # Extension checkbox placed here so it feeds the calculation directly
            st.checkbox("30-day extension granted  [Part V-b-2 / VII]", key="tc_pet2_ext")
            st.caption("Extension requires extenuating circumstances; request within 7 days of results notification.")
            _mail_val = st.session_state.get("tc_pet2_mail")
            _ext_val  = st.session_state.get("tc_pet2_ext", False)
            if _mail_val:
                import datetime as _dt
                _days = 45 + (30 if _ext_val else 0)
                _deadline = _mail_val + _dt.timedelta(days=_days)
                _deadline_str = _deadline.strftime("%m/%d/%Y")
                st.session_state["tc_pet2_deadline"] = _deadline_str
                _label = f"45-Day Deadline{' + 30-Day Extension' if _ext_val else ''}"
                st.markdown(f"**{_label}:** {_deadline_str}  *[Part V-b-2 / VII]*")
            else:
                st.session_state["tc_pet2_deadline"] = ""
                st.markdown("**45-Day Deadline:** *(enter petition mail date)*")
        with col3:
            try:
                pct2 = (float(st.session_state.get("tc_pet2_yes") or 0)
                        / float(st.session_state.get("tc_pet2_total") or 1) * 100)
                pct2_str = f"{pct2:.1f}%"
                st.session_state["tc_pet2_pct"] = pct2_str
                color2 = "green" if pct2 >= 66.7 else "red"
                st.markdown(f"**Support:** :{color2}[{pct2_str}]")
                st.caption("Threshold met" if pct2 >= 66.7 else "Need 66.7% (2/3) to proceed")
            except Exception:
                st.markdown("**Support:** -")
        st.checkbox("Petitions mailed by City on petitioner behalf (City mails twice)  [Part V-b-2 / VII]",
                    key="tc_c_pet2_mailed")
        st.checkbox("Vote eligibility confirmed NOT contingent on HOA membership  [Part V-b-2 / VII]",
                    key="tc_c_pet2_hoa")
        st.checkbox("Non-responses confirmed counted as no votes  [Part V-b-2 / VII]",
                    key="tc_c_pet2_nonresp")
        col1, col2 = st.columns(2)
        with col1:
            tc_date("Last Day of Voting Window  [Part V-b-2 / VII]",
                    key="tc_moratorium_start",
                    help_text="The moratorium clock starts on this date if the petition fails")
        with col2:
            # Moratorium end auto-calculated from last day of voting window.
            # Per Resolution 2026-12 §3: if the vote occurred within 12 months before
            # adoption (02/09/2026), the moratorium is 12 months instead of 24.
            _vote_val = st.session_state.get("tc_moratorium_start")
            if _vote_val:
                import datetime as _dt
                _resolution_date    = _dt.date(2026, 2, 9)
                _twelve_months_prior = _dt.date(2025, 2, 9)
                _short_moratorium = (_twelve_months_prior <= _vote_val <= _resolution_date)
                _months = 12 if _short_moratorium else 24
                # Add months without overflow (handle month-end edge cases)
                _m = _vote_val.month + _months
                _y = _vote_val.year + (_m - 1) // 12
                _m = (_m - 1) % 12 + 1
                import calendar as _cal
                _d = min(_vote_val.day, _cal.monthrange(_y, _m)[1])
                _end = _dt.date(_y, _m, _d)
                _end_str = _end.strftime("%m/%d/%Y")
                st.session_state["tc_moratorium_end"] = _end_str
                if _short_moratorium:
                    st.markdown(f"**Moratorium End:** {_end_str} *(12 months — Res. 2026-12 §3)*")
                    st.caption("Vote occurred within 12 months before adoption — 12-month moratorium applies.")
                else:
                    st.markdown(f"**Moratorium End:** {_end_str} *(24 months)*")
            else:
                st.session_state["tc_moratorium_end"] = ""
                st.markdown("**Moratorium End:** *(enter last day of voting window)*")
        st.caption("Per Resolution 2026-12 Section 3: petitions with votes within 12 months before "
                   "adoption (02/09/2026) use a 12-month moratorium only.")

        if street_class == "Local Residential Street":
            st.markdown("**Cost-Share Agreement - Local Streets [Part VII]**")
            st.checkbox("Petition signed by 2/3 or more of households agreeing to pay 60% of direct costs  [Part VII]",
                        key="tc_c_costshare_agree")
            st.checkbox("HOA funding letter on file confirming HOA will cover cost-share  [Part VII]",
                        key="tc_c_hoa_letter")

    # =========================================================================
    # SECTION VII: COST & PRIORITIZATION
    # =========================================================================
    with st.expander("VII. Cost Estimates & Prioritization Scoring", expanded=False):
        if street_class == "Collector Street":
            st.markdown("**Tier 2 Prioritization Scoring [Part V-b-3 - used when multiple Tier 2 requests are pending]**")
            total_score = 0
            for crit in SCORING_CRITERIA:
                score_val = st.number_input(
                    f"{crit['label']} - max {crit['max']} pts | {crit['basis']}  [{crit['cite']}]",
                    min_value=0, max_value=crit["max"], step=1,
                    key=f"tc_score_{crit['id']}",
                )
                total_score += score_val
            st.session_state["tc_total_score"] = total_score
            st.metric("Total Priority Score", f"{total_score} / 100")
            st.divider()

        st.markdown("**Cost Estimate**")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_input("Estimated Direct Cost ($)  [Part VII]", key="tc_cost_direct")
        with col2:
            try:
                direct  = float(st.session_state.get("tc_cost_direct") or 0)
                conting = direct * 0.10
                total_c = direct + conting
                resident_share = total_c * 0.60
                city_share     = total_c * 0.40
                # Store computed values for the report builder (use setdefault-safe keys
                # that are NOT widget keys, so direct assignment is always safe here)
                st.session_state["tc_cost_contingency"] = f"${conting:,.2f}"
                st.session_state["tc_cost_total"]       = f"${total_c:,.2f}"
                st.session_state["tc_cost_resident"]    = f"${resident_share:,.2f}"
                st.session_state["tc_cost_city"]        = f"${city_share:,.2f}"
                st.markdown(f"**10% Contingency:** ${conting:,.2f}")
                st.markdown(f"**Total Estimated:** ${total_c:,.2f}")
            except Exception:
                resident_share = 0
                city_share = 0
                st.markdown("**Total:** -")
        with col3:
            if street_class == "Local Residential Street":
                st.markdown(f"**Resident Share (60%):** {st.session_state.get('tc_cost_resident', '-')}")
                st.markdown(f"**City Share (40%):** {st.session_state.get('tc_cost_city', '-')}")
                st.caption("Petition expires if 60% payment not received within 6 months of Board approval.")
            elif street_class == "Collector Street":
                st.markdown("**Funding:** 100% City-funded (subject to normal budget process)")

        st.checkbox("Resident 60% payment received prior to installation "
                    "(local streets - must be within 6 months of Board approval)  [Part VII]",
                    key="tc_c_cost_payment")
        st.text_area("Cost / Funding Notes", key="tc_cost_notes", height=100)

    # =========================================================================
    # SECTION VIII: BOARD ACTION
    # =========================================================================
    with st.expander("VIII. Public Meeting & Board of Commissioners Action", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            tc_date("Public Meeting Date  [Part V-b-2]", key="tc_public_meeting_date")
            tc_date("Board Meeting Date  [Part VII]", key="tc_board_date")
        with col2:
            st.text_input("Public Meeting Summary / Outcome", key="tc_public_meeting_notes")
            st.text_input("Board Resolution Number  [Part VII]", key="tc_board_res_num",
                          placeholder="e.g. Resolution 2026-XX")
        st.checkbox("Public meeting scheduled by staff after 2/3 or more petition support received  [Part V-b-2]",
                    key="tc_c_public_meeting")
        st.checkbox("Public meeting conducted; input documented  [Part V-b-2]", key="tc_c_public_conducted")
        st.checkbox("Staff recommendation prepared for Board, incorporating petition results and public meeting input  [Part V-b-2 / VII]",
                    key="tc_c_staff_rec")
        st.checkbox("Board resolution approving proposed locations adopted PRIOR to installation  [Part VII]",
                    key="tc_c_board_res")
        st.checkbox("Board action recorded and outcome documented  [Part VII]", key="tc_c_board_action")
        st.checkbox("Final design complete; City Engineer drainage review done at all locations  [Part VII]",
                    key="tc_c_design_final")
        st.checkbox("Resident 60% payment received before construction begins (expires 6 months after Board approval)  [Part VII]",
                    key="tc_c_payment_rcvd")
        st.checkbox("Improvements installed / constructed", key="tc_c_installed")
        st.checkbox("Complete application file archived (petitions, study, Board resolution, cost records)",
                    key="tc_c_archived")
        st.checkbox("Any leftover funds returned to petitioning group upon project completion  [Part VII]",
                    key="tc_c_leftover_funds")
        st.text_area("Staff Recommendation Summary", key="tc_staff_rec_notes", height=80)
        st.text_area("Final / Closeout Notes", key="tc_final_notes", height=100)

    # =========================================================================
    # EXPORT / CLEAR
    # =========================================================================
    st.divider()
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("Generate Word Document - Checklist + Action Items",
                     type="primary", use_container_width=True, key="btn_tc_export"):
            data = {k: v for k, v in st.session_state.items() if k.startswith("tc_")}
            try:
                buf = build_traffic_calming_report(data)
                street_raw  = st.session_state.get("tc_street_name") or "TC_Review"
                street_safe = (street_raw.replace(" ", "_").replace("/", "-")
                                         .replace("(", "").replace(")", "")[:35])
                filename = f"TC_Review_{street_safe}_{datetime.now(_CT).strftime('%Y%m%d')}.docx"
                st.download_button(
                    label="Download Word Document",
                    data=buf,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="btn_tc_download",
                )
            except Exception as e:
                st.error(f"Error generating document: {e}")
    with col2:
        if st.button("Clear TC Form", use_container_width=True, key="btn_tc_clear"):
            keys_to_clear = [k for k in list(st.session_state.keys()) if k.startswith("tc_")]
            for k in keys_to_clear:
                del st.session_state[k]
            st.rerun()


def main():
    """Main function for Wizard Mode"""
    initialize_session_state()

    page_header(
        title="Engineering Checklist Mode",
        subtitle="Interactive plan review checklist with automatic comment generation",
    )

    # =========================================================================
    # STEP 1: PROJECT SETUP
    # =========================================================================
    st.markdown("<div class='bw-step-heading'>Step 1 — Project Setup</div>", unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Traffic Calming Application is appended at the end of the permit type list.
        # It routes to a separate form (render_traffic_calming_wizard) rather than
        # the standard Yes/No checklist loop used by all other review types.
        ALL_REVIEW_TYPES = REVIEW_TYPES + ["Traffic Calming Application"]

        review_type = st.selectbox(
            "Review Type",
            options=[""] + ALL_REVIEW_TYPES,
            index=0 if not st.session_state.wizard_review_type else (
                ALL_REVIEW_TYPES.index(st.session_state.wizard_review_type) + 1
                if st.session_state.wizard_review_type in ALL_REVIEW_TYPES else 0
            ),
            key="review_type_select"
        )
        
        if review_type and review_type != st.session_state.wizard_review_type:
            st.session_state.wizard_review_type = review_type
            reset_checklist()
            st.rerun()
        elif review_type:
            st.session_state.wizard_review_type = review_type

    # For Traffic Calming, only show Reviewer in col2 — permit number and address
    # are captured inside the TC form itself (case number, street, petitioner).
    is_tc = (st.session_state.wizard_review_type == "Traffic Calming Application")

    with col2:
        if not is_tc:
            permit_number = st.text_input(
                "Permit Number",
                value=st.session_state.wizard_permit_number,
                placeholder="e.g., SW2024-001"
            )
            st.session_state.wizard_permit_number = permit_number
        # col2 intentionally blank for TC (permit info lives in the TC form)

    with col3:
        if not is_tc:
            address = st.text_input(
                "Address",
                value=st.session_state.wizard_address,
                placeholder="e.g., 1808 Sonoma Trce"
            )
            st.session_state.wizard_address = address
        # col3 intentionally blank for TC

    with col4:
        reviewer = st.selectbox(
            "Reviewer",
            options=[""] + REVIEWERS,
            index=0 if not st.session_state.wizard_reviewer else REVIEWERS.index(st.session_state.wizard_reviewer) + 1
        )
        st.session_state.wizard_reviewer = reviewer if reviewer else None
    
    if not st.session_state.wizard_review_type:
        st.markdown('<div class="bw-status-warn">Select a review type above to begin the checklist.</div>', unsafe_allow_html=True)
        return

    # ── Traffic Calming Application routes to its own separate form ──────────
    # All other review types use the standard Yes/No checklist below.
    if st.session_state.wizard_review_type == "Traffic Calming Application":
        if TC_AVAILABLE:
            render_traffic_calming_wizard()
        else:
            st.error(
                "Traffic Calming modules not found in utils/. "
                "Please upload traffic_calming_data.py and traffic_calming_report.py "
                "to the utils/ folder in your GitHub repository."
            )
        return
    
    # =========================================================================
    # STEP 2: INTERACTIVE CHECKLIST
    # Wrapped in @st.fragment so only this section rerenders on button clicks.
    # Sections are st.expanders (collapsed by default) — only the active section
    # renders its items, keeping the widget count low.
    # =========================================================================
    _render_checklist()

    # =========================================================================
    # STANDALONE RESUBMITTAL QUESTION
    # Positioned between the checklist and Step 3, inside its own styled box
    # =========================================================================
    st.markdown('<div class="bw-resubmittal-box">', unsafe_allow_html=True)
    
    resub_col1, resub_col2 = st.columns([3, 1])
    
    with resub_col1:
        st.markdown("**📬 Add standard resubmittal comment?**")
        # Show preview of what BB-0045 says
        resub_text = COMMENTS.get("BB-0045", "")
        if resub_text:
            st.caption(f'BB-0045: "{resub_text}"')
    
    with resub_col2:
        resubmittal = st.radio(
            "Resubmittal",
            options=["—", "Yes", "N/A"],
            index=["—", "Yes", "N/A"].index(st.session_state.wizard_resubmittal) if st.session_state.wizard_resubmittal in ["—", "Yes", "N/A"] else 0,
            key="resubmittal_radio",
            horizontal=True,
            label_visibility="collapsed"
        )
        st.session_state.wizard_resubmittal = resubmittal
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # =========================================================================
    # STEP 3: REVIEW SUMMARY & EXPORT
    # =========================================================================
    st.markdown("---")
    st.markdown("<div class='bw-step-heading'>📊 Step 3 — Review Summary &amp; Export</div>", unsafe_allow_html=True)
    
    yes_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "Yes")
    no_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "No")
    na_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "N/A")
    
    # Include resubmittal in the "has comments" logic
    has_comments = no_count > 0 or st.session_state.wizard_resubmittal == "Yes"
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("✅ Compliant", yes_count)
    with col2:
        st.metric("❌ Issues Found", no_count)
    with col3:
        st.metric("➖ N/A", na_count)
    with col4:
        st.metric("📝 Total Reviewed", yes_count + no_count + na_count)
    
    # Export section
    st.markdown('<div class="bw-export-section">', unsafe_allow_html=True)
    st.markdown("### 📤 Export Review")
    
    if not has_comments and (yes_count + na_count) > 0:
        st.success("✅ No issues found! All reviewed items are compliant.")
    elif has_comments:
        comment_parts = []
        if no_count > 0:
            comment_parts.append(f"{no_count} issue(s) found")
        if st.session_state.wizard_resubmittal == "Yes":
            comment_parts.append("resubmittal comment included")
        st.warning(f"⚠️ {' + '.join(comment_parts)}.")
    
    # Row 1: Word Document + Clear Review
    col1, col2 = st.columns(2)
    
    with col1:
        if DOCX_AVAILABLE:
            if st.button("📄 Generate Word Document", type="primary", use_container_width=True):
                if not st.session_state.wizard_permit_number:
                    st.error("Please enter a permit number before exporting.")
                elif len(st.session_state.wizard_checklist_state) == 0:
                    st.error("Please review at least one item before exporting.")
                else:
                    doc_buffer = generate_word_document()
                    if doc_buffer:
                        filename = f"Review_{st.session_state.wizard_permit_number}_{datetime.now(_CT).strftime('%Y%m%d')}.docx"
                        st.download_button(
                            label="⬇️ Download Word Document",
                            data=doc_buffer,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
        else:
            st.warning("python-docx not available. Install it to enable Word export.")
    
    with col2:
        if st.button("🗑️ Clear Review", use_container_width=True):
            reset_checklist()
            st.session_state.wizard_permit_number = ""
            st.session_state.wizard_address = ""
            st.session_state.wizard_reviewer = None
            st.rerun()

    # Row 2: LAMA CSV + Bluebeam BAX (shown when there are any comments)
    if has_comments:
        st.markdown("#### 📊 Extract Comments")

        permit_num = st.session_state.wizard_permit_number or "review"
        datestamp = datetime.now(_CT).strftime('%Y%m%d')

        col_e1, col_e2 = st.columns(2)

        with col_e1:
            lama_data = generate_lama_csv()
            if lama_data:
                st.download_button(
                    label="📥 Create CSV File of Comments",
                    data=lama_data,
                    file_name=f"LAMA_Comments_{permit_num}_{datestamp}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    help="Single-column CSV for the LAMA Comment Uploader extension"
                )

        with col_e2:
            bax_data = generate_bluebeam_bax()
            if bax_data:
                st.download_button(
                    label="📐 Create Bluebeam Comments File",
                    data=bax_data,
                    file_name=f"Markups_{permit_num}_{datestamp}.bax",
                    mime="application/octet-stream",
                    use_container_width=True,
                    help="Import into Bluebeam via Markup → Import (.bax format with full styling)"
                )
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Quick copy section for comments
    if has_comments:
        st.markdown("---")
        st.subheader("📋 Quick Copy - All Comments")
        st.caption("Copy these comments directly into Bluebeam or your permit system:")
        
        all_comments = []
        checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
        for section_id, section_data in checklist.items():
            for item in section_data["items"]:
                item_key = item["id"]
                if st.session_state.wizard_checklist_state.get(item_key) == "No":
                    selected = st.session_state.wizard_selected_comments.get(item_key, [])
                    custom_note = st.session_state.wizard_custom_notes.get(item_key, "")
                    
                    for comment_id in selected:
                        comment_text = COMMENTS.get(comment_id, "")
                        if comment_text:
                            all_comments.append(f"[{comment_id}] {comment_text}")
                    
                    if custom_note.strip():
                        all_comments.append(f"[CUSTOM] {custom_note}")
        
        # Append resubmittal at the end
        if st.session_state.wizard_resubmittal == "Yes":
            resub_text = COMMENTS.get("BB-0045", "")
            if resub_text:
                all_comments.append(f"[BB-0045] {resub_text}")
        
        if all_comments:
            comments_text = "\n\n".join(f"{i+1}. {c}" for i, c in enumerate(all_comments))
            st.text_area(
                "All Comments",
                value=comments_text,
                height=300,
                label_visibility="collapsed"
            )
    
    # Navigation
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🏠 Home"):
            st.switch_page("app.py")
    with col2:
        if st.button("💬 Q&A Mode"):
            st.switch_page("pages/1_QA_Mode.py")


if __name__ == "__main__":
    main()
    footer()
