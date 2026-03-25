"""
==============================================================================
TRAFFIC CALMING REPORT — Word Document Builder
==============================================================================
File: utils/traffic_calming_report.py
Purpose: Generates the Traffic Calming Application Review .docx report
         from the session state data collected in the Wizard Mode form.

Uses python-docx (already in requirements.txt).
Policy Authority: Resolution 2026-12, City of Brentwood, TN (02/09/2026)

Called by: pages/2_Wizard_Mode.py → render_traffic_calming_wizard()
==============================================================================
"""

from io import BytesIO
from datetime import datetime


def build_traffic_calming_report(data: dict) -> BytesIO:
    """
    Build a .docx report from traffic calming form data.

    Args:
        data (dict): All session_state values whose keys start with 'tc_'
    Returns:
        BytesIO: Buffer containing the completed .docx file ready for download
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Brand colors
    NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
    ORANGE = RGBColor(0xE8, 0x65, 0x1A)
    GREEN  = RGBColor(0x2D, 0x6A, 0x4F)
    GREY   = RGBColor(0x8A, 0x82, 0x78)

    doc = Document()

    # ── Page setup: US Letter, 1-inch margins ─────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.left_margin  = sec.right_margin  = Inches(1.0)
    sec.top_margin   = sec.bottom_margin = Inches(1.0)

    # ── Convenience helpers ────────────────────────────────────────────────

    def h1(text):
        """Top-level section heading — navy, 14pt bold."""
        p = doc.add_heading("", level=1)
        p.clear()
        run = p.add_run(text)
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def h2(text):
        """Sub-heading — navy, 11pt bold."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def kv(label, value, cite=""):
        """
        Key-value paragraph.
        Bold label: normal value  [italic cite in navy]
        """
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r2 = p.add_run(str(value) if value else "—")
        r2.font.size = Pt(9.5)
        if cite:
            r3 = p.add_run(f"  [{cite}]")
            r3.italic = True
            r3.font.size = Pt(8)
            r3.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(1)
        return p

    def checked(label, is_checked, cite=""):
        """Checkbox line: ☑/☐ label [cite]"""
        mark = "☑" if is_checked else "☐"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(f"{mark}  {label}")
        run.font.size = Pt(9.5)
        if not is_checked:
            run.font.color.rgb = RGBColor(0x99, 0x44, 0x00)   # amber for unchecked
        if cite:
            r2 = p.add_run(f"  [{cite}]")
            r2.italic = True
            r2.font.size = Pt(8)
            r2.font.color.rgb = NAVY
        return p

    def divider():
        """Thin navy horizontal rule."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:color'), '1B3A6B')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def get(key):
        """Safe getter from data dict."""
        return data.get(key) or ""

    def is_checked(key):
        """Return bool from data dict."""
        return bool(data.get(key))

    # Derive key values used throughout
    street_class = get("tc_street_class")
    is_collector = street_class == "Collector Street"
    is_local     = street_class == "Local Residential Street"
    is_arterial  = street_class == "Arterial Street"

    # ── TITLE BLOCK ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("TRAFFIC CALMING APPLICATION REVIEW")
    tr.font.size  = Pt(16)
    tr.font.bold  = True
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "City of Brentwood, Tennessee  ·  Resolution 2026-12  ·  Engineering Department"
    )
    sr.font.size  = Pt(9)
    sr.font.color.rgb = ORANGE

    # Third line: case number + street name (only if at least one is populated)
    case_num   = get("tc_case_num")
    street_nm  = get("tc_street_name")
    if case_num or street_nm:
        id_line = doc.add_paragraph()
        id_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = [p for p in [case_num, street_nm] if p]
        id_run = id_line.add_run("  —  ".join(parts))
        id_run.font.size  = Pt(11)
        id_run.font.bold  = True
        id_run.font.color.rgb = NAVY

    divider()

    # ── SECTION I: ADMINISTRATIVE ─────────────────────────────────────────
    h1("I. Administrative Review")
    kv("Case / File Number",    get("tc_case_num"))
    kv("Application Date",      get("tc_app_date_str") or get("tc_app_date"))
    kv("Street Name",           get("tc_street_name"),    "Municipal Code — Street Classification")
    kv("Street Segment / Limits", get("tc_street_segment"))
    kv("Street Classification", get("tc_street_class"),   "Part V / VII")
    kv("Application Type",      get("tc_app_type"),       "Parts V, VII")
    kv("Petitioner",            get("tc_petitioner_name"))
    kv("Petitioner Contact",    get("tc_petitioner_contact"))
    kv("HOA Determination",     get("tc_hoa_status"),     "Part V / VII — 90-day rule")
    kv("Eligible Residences",   get("tc_eligible_res"),   "Part V / VII — within 600 ft")
    kv("Signatures Received",   get("tc_sigs_received"))
    kv("Initial Support %",     get("tc_init_pct"),       "Part V / VII (>50% required)")
    kv("Stated Problem",        get("tc_problem_desc"))

    checked("HOA governance confirmed",                             is_checked("tc_c_hoa_gov"),    "Part V")
    checked("HOA written request submitted (or bypass justified)",  is_checked("tc_c_hoa_req"),    "Part V / VII — 90-day rule")
    checked("Initial >50% support petition obtained",               is_checked("tc_c_init_pet"),   "Part V / VII")
    checked("One signature and printed name per residence",         is_checked("tc_c_pet_format"), "Part V")

    divider()

    # ── SECTION II: CLASSIFICATION ────────────────────────────────────────
    h1("II. Street Classification & Eligibility")
    if is_collector:
        kv("Funding",  "100% City-funded (subject to budget approval)", "Part V – Funding")
        kv("Policy",   "Traffic Calming Policy — Part V")
    elif is_local:
        kv("Funding",  "Residents pay 60% of direct costs + 10% contingency", "Part VII")
        kv("Policy",   "Speed Hump Policy — Part VII")
    elif is_arterial:
        kv("Note",     "Arterial — standard traffic calming policy does NOT apply. Engineering study required.", "Part V")

    checked("Public street confirmed (not gated/private)",                    is_checked("tc_c_not_private"),    "Part V / VII")
    checked("Not a designated primary emergency route",                        is_checked("tc_c_not_emergency"),  "Part V")
    checked("Collector: verified on City identified residential collector list",is_checked("tc_c_collector_list"),"Part V – Introduction")

    divider()

    # ── SECTION III: DATA COLLECTION ─────────────────────────────────────
    h1("III. Data Collection & Field Review")
    kv("85th Percentile Speed",    f"{get('tc_data_85th')} mph",   "Part V–a / Part VII")
    kv("Posted Speed Limit",       f"{get('tc_data_limit')} mph")
    kv("Speed Excess over Limit",  get("tc_speed_excess"),         "Part V–a (≥8 mph required for collectors)")
    kv("Average Daily Traffic",    f"{get('tc_data_adt')} vpd",    "Part V–a / Part VII")
    kv("Crash Count (12 months)",  get("tc_data_crashes"),         "Parts I–II / Part V–b–1")
    kv("Cut-Through Traffic %",    f"{get('tc_data_cutthru')}%",   "Part VII (≥35% = cut-through condition)")
    kv("School Walking Route",     get("tc_data_school_route"),    "Part V–b–3")
    kv("Continuous Sidewalk",      get("tc_data_sidewalk_status"), "Part V–b–3")
    kv("Data Collection Notes",    get("tc_data_notes"))

    checked("Speed study completed (≥24-hr weekday for collectors)", is_checked("tc_c_data_speed"),   "Part V–a")
    checked("Traffic count / ADT collected",                          is_checked("tc_c_data_vol"),     "Part V–a / VII")
    checked("Crash history reviewed (12 months minimum)",             is_checked("tc_c_data_crash"),   "Parts I–II")
    checked("School route status confirmed",                          is_checked("tc_c_data_school"),  "Part V–b–3")
    checked("Continuous sidewalk presence/absence confirmed",         is_checked("tc_c_data_sidewalk"),"Part V–b–3")

    divider()

    # ── SECTION IV: STREET-TYPE SPECIFIC ─────────────────────────────────
    if is_collector:
        h1("IV. Traffic Calming — Collector Street Criteria (Part V)")
        h2("Roadway Eligibility Criteria — All must be met (Part V–a)")

        # Auto-evaluate criteria
        try:
            adt_ok  = float(get("tc_data_adt") or 0) >= 500
            len_ok  = float(get("tc_seg_length") or 0) >= 800
            spd_ok  = float(get("tc_speed_excess_raw") or 0) >= 8
        except:
            adt_ok = len_ok = spd_ok = False

        kv("ADT ≥ 500/day",             "✓ MET" if adt_ok  else "✗ NOT MET", "Part V–a – Volume")
        kv("Segment ≥ 800 ft",          "✓ MET" if len_ok  else "✗ NOT MET", "Part V–a – Other Criteria")
        kv("85th %ile speed ≥8 mph over limit", "✓ MET" if spd_ok else "✗ NOT MET", "Part V–a – Speed")

        checked("24-hr weekday speed data collected",                    is_checked("tc_c_coll_speed_data"),    "Part V–a")
        checked("Street has ≤2 traffic lanes",                           is_checked("tc_c_coll_2lanes"),        "Part V–a")
        checked("Logical termini for calming treatment identifiable",    is_checked("tc_c_coll_termini"),       "Part V–a")

        h2("Engineering Study Contents Required (Part V–b–1)")
        checked("Study includes traffic volume analysis",                is_checked("tc_c_coll_study_vol"),     "Part V–b–1")
        checked("Study includes traffic speed analysis",                 is_checked("tc_c_coll_study_speed"),   "Part V–b–1")
        checked("Study includes accident history review",                is_checked("tc_c_coll_study_crash"),   "Part V–b–1")
        checked("Study notes sidewalk presence/absence",                 is_checked("tc_c_coll_study_sidewalk"),"Part V–b–1")
        checked("Study addresses school walking route",                  is_checked("tc_c_coll_study_school"),  "Part V–b–1")
        checked("Study outlines applicable Tier 2 strategies",           is_checked("tc_c_coll_study_t2"),      "Part V–b–1")

    elif is_local:
        h1("IV. Speed Hump Eligibility — Local Street (Part VII)")
        h2("Street Eligibility Criteria — All must be met (Part VII)")
        kv("Street Width",          f"{get('tc_local_width')} ft",      "Part VII (<30 ft required)")
        kv("Street Grade",          f"{get('tc_local_grade')}%",        "Part VII (≤6% required)")
        kv("Posted Speed Limit",    f"{get('tc_local_spd_limit')} mph", "Part VII (≤30 mph required)")
        kv("Current ADT",           f"{get('tc_local_adt')} vpd",       "Part VII (500–2,500 required)")
        kv("Projected Full-Build ADT", f"{get('tc_local_adt_proj')} vpd","Part VII (≤2,500 at full build)")

        checked("Two-lane street <30 ft wide confirmed",                  is_checked("tc_c_local_lanewidth"),   "Part VII")
        checked("Grade ≤6% confirmed",                                    is_checked("tc_c_local_grade"),       "Part VII")
        checked("Speed limit ≤30 mph confirmed",                          is_checked("tc_c_local_speedlimit"),  "Part VII")
        checked("Not an arterial or collector street",                    is_checked("tc_c_local_notarterial"), "Part VII")
        checked("Cut-through or speeding problem identified",             is_checked("tc_c_local_cutthru"),     "Part VII")
        checked("Connection route between arterial/collector or subdivision pass-through", is_checked("tc_c_local_connection"), "Part VII")

        h2("Design Requirements (Part VII)")
        checked("Minimum 2 humps; spacing 300–600 ft apart",             is_checked("tc_c_hump_spacing"),      "Part VII")
        checked("200-ft clearance from intersections and sharp curves",   is_checked("tc_c_hump_clearance"),    "Part VII")
        checked("Height 3–4 inches; 12-ft travel length per standard",   is_checked("tc_c_hump_dims"),         "Part VII")
        checked("Regulatory 'Speed Control District' signs installed",    is_checked("tc_c_hump_signage"),      "Part VII")
        checked("MUTCD advance warning signs + 15 MPH advisory plate",   is_checked("tc_c_hump_warn"),         "Part VII")
        checked("Double yellow centerline + 12-in white stripe markings",is_checked("tc_c_hump_markings"),     "Part VII – Pavement Marking Details")
        checked("City Engineer drainage review completed",                is_checked("tc_c_hump_drainage"),     "Part VII")

    divider()

    # ── SECTION V: TIER 1 ────────────────────────────────────────────────
    h1("V. Tier 1 — Non-Construction Strategies (Part V–b–1)")
    kv("Tier 1 Implementation Date",  get("tc_t1_date_str") or get("tc_t1_date"),        "Part V–b–1")
    kv("Six-Month Review Date",       get("tc_t1_review_date_str") or get("tc_t1_review_date"), "Part V–b–1 (effectiveness review)")
    tier1_outcome = (
        "Effective — no further action required" if is_checked("tc_c_t1_effective")
        else "Ineffective after 6 months — Tier 2 requested" if is_checked("tc_c_t1_ineffective")
        else "Pending evaluation"
    )
    kv("Tier 1 Outcome", tier1_outcome, "Part V–b–1")
    kv("Tier 1 Notes",   get("tc_t1_notes"))

    checked("Study recommendation includes one or more Tier 1 strategies", is_checked("tc_c_t1_study"),       "Part V–b–1")
    checked("Staff met with petitioner to outline study recommendations",   is_checked("tc_c_t1_petitioner"),  "Part V–b–1")
    checked("Tier 1 strategy implemented",                                  is_checked("tc_c_t1_implemented"), "Part V–b–1")

    divider()

    # ── SECTION VI: TIER 2 / SECOND PETITION ─────────────────────────────
    h1("VI. Tier 2 / Speed Humps — Second Petition (Parts V–b–2; VII)")

    # Selected strategies
    t2_selected = data.get("tc_t2_strategies") or []
    if t2_selected:
        h2("Tier 2 Strategies Proposed")
        for s in t2_selected:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(f"• {s}").font.size = Pt(9.5)

    h2("Second-Round Petition Results")
    kv("Total Households in Affected Area", get("tc_pet2_total"),    "Part V–b–2 / VII (within 600 ft)")
    kv("Yes Votes Received",                get("tc_pet2_yes"),      "Part V–b–2 / VII")
    kv("Support Percentage",                get("tc_pet2_pct"),      "≥66.7% (2/3) required")
    kv("Petition Mail Date",                get("tc_pet2_mail_str") or get("tc_pet2_mail"),     "Part V–b–2 / VII")
    kv("45-Day Response Deadline",          get("tc_pet2_deadline"), "Part V–b–2 / VII")
    kv("30-Day Extension Granted",          "Yes" if is_checked("tc_pet2_ext") else "No", "Part V–b–2 / VII")
    kv("Last Day of Voting Window",         get("tc_moratorium_start_str") or get("tc_moratorium_start"), "Part V–b–2 / VII")
    kv("Moratorium End Date (if failed)",   get("tc_moratorium_end"),"Part V–b–2; Res. 2026-12 §3")

    checked("Second study confirms Tier 1 was ineffective",         is_checked("tc_c_t2_validate"),    "Part V–b–2")
    checked("Traffic Engineer reviewed Tier 2 recommendation",      is_checked("tc_c_t2_trafficeng"),  "Part V–b–2")
    checked("Separate petition prepared for each improvement",       is_checked("tc_c_t2_sep_petition"),"Part V–b–2")
    checked("Petitions mailed by City (sent twice)",                 is_checked("tc_c_pet2_mailed"),    "Part V–b–2 / VII")
    checked("Vote eligibility not tied to HOA membership status",    is_checked("tc_c_pet2_hoa"),       "Part V–b–2 / VII")
    checked("Non-responses counted as 'no' votes",                   is_checked("tc_c_pet2_nonresp"),   "Part V–b–2 / VII")

    if is_local:
        h2("Cost-Share Agreement (Local Streets — Part VII)")
        checked("≥2/3 households agreed to pay 60% of direct costs", is_checked("tc_c_costshare_agree"),"Part VII")
        checked("HOA funding letter on file (if HOA is paying)",      is_checked("tc_c_hoa_letter"),     "Part VII")

    divider()

    # ── SECTION VII: COST & PRIORITIZATION ───────────────────────────────
    h1("VII. Cost Estimates & Prioritization Scoring")

    if is_collector:
        h2("Tier 2 Prioritization Score (Part V–b–3)")
        from traffic_calming_data import SCORING_CRITERIA
        for crit in SCORING_CRITERIA:
            score_val = data.get(f"tc_score_{crit['id']}", 0) or 0
            kv(f"{crit['label']} (max {crit['max']})", f"{score_val} pts — {crit['basis']}", crit["cite"])
        kv("TOTAL PRIORITY SCORE", f"{get('tc_total_score')} / 100", "Part V–b–3")
        doc.add_paragraph()

    kv("Estimated Direct Cost",          get("tc_cost_direct"),      "Part VII")
    kv("10% Contingency",                get("tc_cost_contingency"), "Part VII")
    kv("Total Estimated Cost",           get("tc_cost_total"),       "Part VII")
    if is_local:
        kv("Resident Share (60%)",       get("tc_cost_resident"),    "Part VII")
        kv("City Share (40%)",           get("tc_cost_city"),        "Part VII")
        checked("Resident 60% payment received prior to installation (expires 6 months post-Board approval)",
                is_checked("tc_c_cost_payment"), "Part VII")
    else:
        kv("Funding",                    "100% City-funded (subject to normal budgeting process)", "Part V – Funding")
    kv("Cost / Funding Notes",           get("tc_cost_notes"))

    divider()

    # ── SECTION VIII: BOARD ACTION ────────────────────────────────────────
    h1("VIII. Public Meeting & Board of Commissioners Action")
    kv("Public Meeting Date",    get("tc_public_meeting_date_str") or get("tc_public_meeting_date"),  "Part V–b–2")
    kv("Meeting Outcome",        get("tc_public_meeting_notes"))
    kv("Board Meeting Date",     get("tc_board_date_str") or get("tc_board_date"),           "Part VII")
    kv("Board Resolution No.",   get("tc_board_res_num"),        "Part VII")
    kv("Staff Recommendation",   get("tc_staff_rec_notes"),      "Part V–b–2 / VII")
    kv("Final / Closeout Notes", get("tc_final_notes"))

    checked("Public meeting scheduled by staff after ≥2/3 petition support received", is_checked("tc_c_public_meeting"),  "Part V–b–2")
    checked("Public meeting conducted; input documented",                               is_checked("tc_c_public_conducted"),"Part V–b–2")
    checked("Staff recommendation prepared for Board",                                  is_checked("tc_c_staff_rec"),       "Part V–b–2 / VII")
    checked("Board resolution adopted prior to installation",                           is_checked("tc_c_board_res"),       "Part VII")
    checked("Board action recorded",                                                    is_checked("tc_c_board_action"),    "Part VII")
    checked("Final design complete; City Engineer drainage review done",                is_checked("tc_c_design_final"),    "Part VII")
    checked("60% resident payment received before installation",                        is_checked("tc_c_payment_rcvd"),    "Part VII")
    checked("Improvements installed",                                                   is_checked("tc_c_installed"),       "")
    checked("Complete application file archived",                                       is_checked("tc_c_archived"),        "")
    checked("Leftover funds returned to petitioning group",                             is_checked("tc_c_leftover_funds"),  "Part VII")

    divider()

    # ── SECTION IX: ACTION ITEMS ──────────────────────────────────────────
    h1("IX. Outstanding Action Items")

    # Master list: (session_key, display_label, policy_cite, section_label, applies_to)
    # applies_to: "all" | "local" | "collector"
    # Items marked "local" only appear for Local Residential Street applications.
    # Items marked "collector" only appear for Collector Street applications.
    ALL_ACTIONS = [
        # Admin — applies to all
        ("tc_c_hoa_gov",          "Confirm HOA governance status",                               "Part V",         "I. Admin",            "all"),
        ("tc_c_hoa_req",          "Confirm HOA written request submitted or bypass justified",    "Part V / VII",   "I. Admin",            "all"),
        ("tc_c_init_pet",         "Obtain initial >50% support petition",                        "Part V / VII",   "I. Admin",            "all"),
        ("tc_c_pet_format",       "Verify one signature/printed name per residence",              "Part V",         "I. Admin",            "all"),
        # Classification — applies to all
        ("tc_c_not_private",      "Confirm public street (not gated/private)",                   "Part V / VII",   "II. Classification",  "all"),
        ("tc_c_not_emergency",    "Confirm not a designated primary emergency route",             "Part V",         "II. Classification",  "all"),
        ("tc_c_collector_list",   "Verify collector street list status",                         "Part V",         "II. Classification",  "collector"),
        # Data — applies to all
        ("tc_c_data_speed",       "Complete speed study (≥24-hr weekday)",                       "Part V–a",       "III. Data",           "all"),
        ("tc_c_data_vol",         "Complete traffic count / ADT",                                "Part V–a / VII", "III. Data",           "all"),
        ("tc_c_data_crash",       "Review crash history (12 months)",                            "Parts I–II",     "III. Data",           "all"),
        ("tc_c_data_school",      "Confirm school route status",                                 "Part V–b–3",     "III. Data",           "all"),
        ("tc_c_data_sidewalk",    "Confirm continuous sidewalk status",                          "Part V–b–3",     "III. Data",           "all"),
        # Collector-specific criteria — only for Collector Street
        ("tc_c_coll_speed_data",  "Collect ≥24-hr weekday speed data",                          "Part V–a",       "IV. Collector",       "collector"),
        ("tc_c_coll_2lanes",      "Confirm street has ≤2 traffic lanes",                         "Part V–a",       "IV. Collector",       "collector"),
        ("tc_c_coll_termini",     "Identify logical termini for calming treatment",              "Part V–a",       "IV. Collector",       "collector"),
        ("tc_c_coll_study_vol",   "Include volume analysis in engineering study",                "Part V–b–1",     "IV. Collector",       "collector"),
        ("tc_c_coll_study_speed", "Include speed analysis in engineering study",                 "Part V–b–1",     "IV. Collector",       "collector"),
        ("tc_c_coll_study_crash", "Include accident history in engineering study",               "Part V–b–1",     "IV. Collector",       "collector"),
        ("tc_c_coll_study_sidewalk","Note sidewalk status in engineering study",                 "Part V–b–1",     "IV. Collector",       "collector"),
        ("tc_c_coll_study_school","Address school walking route in engineering study",            "Part V–b–1",     "IV. Collector",       "collector"),
        ("tc_c_coll_study_t2",    "Outline applicable Tier 2 strategies in study",               "Part V–b–1",     "IV. Collector",       "collector"),
        # Local-specific criteria — only for Local Residential Street
        ("tc_c_local_lanewidth",  "Confirm street is <30 ft wide (two lanes)",                  "Part VII",       "IV. Local",           "local"),
        ("tc_c_local_grade",      "Confirm street grade ≤6%",                                   "Part VII",       "IV. Local",           "local"),
        ("tc_c_local_speedlimit", "Confirm posted speed limit ≤30 mph",                         "Part VII",       "IV. Local",           "local"),
        ("tc_c_local_notarterial","Confirm street is not arterial or collector",                 "Part VII",       "IV. Local",           "local"),
        ("tc_c_local_cutthru",    "Identify cut-through or speeding problem (with data)",        "Part VII",       "IV. Local",           "local"),
        ("tc_c_local_connection", "Confirm connection route or subdivision pass-through",        "Part VII",       "IV. Local",           "local"),
        ("tc_c_hump_spacing",     "Verify minimum 2 humps; spacing 300–600 ft",                 "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_clearance",   "Verify 200-ft clearance from intersections and curves",      "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_dims",        "Verify hump height 3–4 in; 12-ft travel length",             "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_signage",     "Install regulatory 'Speed Control District' signs",          "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_warn",        "Install MUTCD advance warning signs and 15 MPH plate",       "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_markings",    "Install pavement markings per standard details",             "Part VII",       "IV. Local – Design",  "local"),
        ("tc_c_hump_drainage",    "City Engineer reviews drainage at all hump locations",        "Part VII",       "IV. Local – Design",  "local"),
        # Tier 1 — applies to all
        ("tc_c_t1_study",         "Include Tier 1 strategy in study recommendation",            "Part V–b–1",     "V. Tier 1",           "all"),
        ("tc_c_t1_petitioner",    "Staff meets with petitioner on study findings",              "Part V–b–1",     "V. Tier 1",           "all"),
        ("tc_c_t1_implemented",   "Implement Tier 1 strategy",                                 "Part V–b–1",     "V. Tier 1",           "all"),
        # Tier 2 — applies to all
        ("tc_c_t2_validate",      "Conduct second study confirming Tier 1 ineffective",         "Part V–b–2",     "VI. Tier 2",          "all"),
        ("tc_c_t2_trafficeng",    "Traffic Engineer reviews Tier 2 recommendation",             "Part V–b–2",     "VI. Tier 2",          "all"),
        ("tc_c_t2_sep_petition",  "Prepare separate petition for each improvement",             "Part V–b–2",     "VI. Tier 2",          "all"),
        ("tc_c_pet2_mailed",      "City mails second-round petitions (mailed twice)",           "Part V–b–2 / VII","VI. Tier 2",         "all"),
        ("tc_c_pet2_hoa",         "Confirm votes not tied to HOA membership",                   "Part V–b–2 / VII","VI. Tier 2",         "all"),
        ("tc_c_pet2_nonresp",     "Confirm non-responses counted as no votes",                  "Part V–b–2 / VII","VI. Tier 2",         "all"),
        ("tc_c_costshare_agree",  "Obtain cost-share agreement (60%) from ≥2/3 households",    "Part VII",       "VI. Tier 2",          "local"),
        ("tc_c_hoa_letter",       "Obtain HOA funding letter (if HOA is paying)",               "Part VII",       "VI. Tier 2",          "local"),
        # Board — applies to all
        ("tc_c_public_meeting",   "Schedule public meeting",                                    "Part V–b–2",     "VIII. Board",         "all"),
        ("tc_c_public_conducted", "Conduct public meeting; document input",                     "Part V–b–2",     "VIII. Board",         "all"),
        ("tc_c_staff_rec",        "Prepare staff recommendation for Board",                     "Part V–b–2 / VII","VIII. Board",        "all"),
        ("tc_c_board_res",        "Board adopts resolution approving location(s)",               "Part VII",       "VIII. Board",         "all"),
        ("tc_c_board_action",     "Record Board action / outcome",                              "Part VII",       "VIII. Board",         "all"),
        ("tc_c_design_final",     "Complete final design; City Engineer drainage review",       "Part VII",       "VIII. Board",         "all"),
        ("tc_c_payment_rcvd",     "Receive 60% resident payment before installation",           "Part VII",       "VIII. Board",         "local"),
        ("tc_c_installed",        "Improvements installed / constructed",                       "",               "VIII. Board",         "all"),
        ("tc_c_archived",         "Complete application file archived",                         "",               "VIII. Board",         "all"),
        ("tc_c_leftover_funds",   "Return leftover funds to petitioning group",                 "Part VII",       "VIII. Board",         "all"),
    ]

    # Filter to only the actions relevant to this application's classification.
    # "all" items always appear. "local" items only for Local Residential Street.
    # "collector" items only for Collector Street.
    def applies(classification_tag):
        if classification_tag == "all":
            return True
        if classification_tag == "local" and is_local:
            return True
        if classification_tag == "collector" and is_collector:
            return True
        return False

    relevant = [(k, lbl, cite, sec) for k, lbl, cite, sec, cls in ALL_ACTIONS if applies(cls)]
    pending  = [(k, lbl, cite, sec) for k, lbl, cite, sec in relevant if not data.get(k)]
    done     = [(k, lbl, cite, sec) for k, lbl, cite, sec in relevant if data.get(k)]

    # Completed items
    h2(f"Completed Steps ({len(done)} items)")
    if done:
        for _, lbl, cite, sec in done:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"☑  [{sec}]  {lbl}")
            run.font.size = Pt(9)
            run.font.color.rgb = GREEN
            if cite:
                r2 = p.add_run(f"  [{cite}]")
                r2.italic = True
                r2.font.size = Pt(8)
                r2.font.color.rgb = NAVY
    else:
        doc.add_paragraph("No items completed yet.").runs[0].font.color.rgb = GREY

    doc.add_paragraph()

    # Outstanding action items (numbered, bold)
    h2(f"Outstanding Action Items — {len(pending)} remaining (in process order)")
    if pending:
        for i, (_, lbl, cite, sec) in enumerate(pending, 1):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"[{sec}]  {lbl}")
            run.font.size = Pt(9.5)
            run.font.bold = True
            if cite:
                r2 = p.add_run(f"  [{cite}]")
                r2.italic = True
                r2.font.bold = False
                r2.font.size = Pt(8)
                r2.font.color.rgb = NAVY
    else:
        p = doc.add_paragraph("All action items complete. Ready for final documentation.")
        for run in p.runs:
            run.font.color.rgb = GREEN
            run.font.bold = True

    divider()

    # ── FOOTER ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    tz_label = ""
    try:
        import pytz
        cst = pytz.timezone("America/Chicago")
        now_cst = datetime.now(cst)
        timestamp = now_cst.strftime("%B %d, %Y  %I:%M %p CST")
    except:
        timestamp = datetime.now().strftime("%B %d, %Y  %I:%M %p")

    foot = doc.add_paragraph(
        f"Report generated: {timestamp}  |  "
        "Policy: Resolution 2026-12, City of Brentwood, TN (adopted 02/09/2026)  |  "
        "Engineering Department"
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in foot.runs:
        run.font.size   = Pt(7.5)
        run.italic      = True
        run.font.color.rgb = GREY

    # ── Return buffer ──────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
